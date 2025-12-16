#!/usr/bin/env python3
# -*- coding: utf-8 -*- 
"""
@Project : tools
@File    : test_record.py 
@Author  : Shawn
@Date    : 2025/12/16 20:38 
@Info    : test合集，知识记录
"""

import datetime


def get_datetime():
    time = datetime.datetime.now().strftime('%y%m%d%H%M%S')
    return time


def test1():
    """excel形状"""
    import copy
    import datetime
    import io
    import os
    import time
    from pathlib import Path
    from typing import Dict, List, Optional

    import openpyxl
    import pythoncom
    import win32com
    from PIL import Image
    from openpyxl import load_workbook
    from openpyxl.cell.cell import Cell, MergedCell
    from openpyxl.drawing.image import Image as XLImage
    from openpyxl.styles import Font, PatternFill, Border, Side, Alignment, Protection

    import win32com.client as win32
    import win32com
    print(win32com.__gen_path__)
    from PIL import ImageGrab
    import time

    # # 启动Excel应用
    excel = win32.gencache.EnsureDispatch('Excel.Application')
    # excel.Visible = True  # 设置为True以便观察，False则后台运行

    input_file = r"D:\Code\Project\tools\data\test\4.xlsx"
    output_folder = r"D:\Code\Project\tools\data\temp\pic"
    output_file = r"D:\Code\Project\tools\data\temp\test251205-xxxx.xlsx"

    img_path = r"D:\Code\Project\tools\data\temp\pic\Group 15_251205150208_13.png"

    # excel = openpyxl.Workbook()
    # # workbook = excel.Workbooks.Open(input_file)
    # # 清除已经存在的sheet
    # default_sheet = excel.active
    # excel.remove(default_sheet)
    # sheet = excel.create_sheet("test123")
    # try:
    #     img = XLImage(img_path)
    #     img.width, img.height = (317.47, 177.80)
    #     sheet.add_image(img)
    #     print(f"{sheet}插入图片成功")
    #
    #     excel.save(output_file)
    #     print(f"{output_file}已保存成功！")
    #
    # except Exception as e:
    #     print(e)
    # excelApp=None
    excelApp = win32.Dispatch("Excel.Application")
    excelApp.Visible = False

    excelApp.Quit()

    workbook = excelApp.Workbooks.Open(input_file)
    sheet = workbook.Worksheets("M0901")

    shape = sheet.Shapes("TextBox 7")
    new_sheet = workbook.Worksheets("test")
    shape.Copy()
    new_sheet.Paste()

    # img = sheet.Pictures().Insert(img_path)
    # img.Left = 96.79
    # img.Top = 467.02
    # img.Width = 317.47
    # img.Height = 177.80

    # 复制sheet
    # sheet.Copy(pythoncom.Empty, workbook.Sheets(workbook.Sheets.Count))
    # new_sheet = workbook.Worksheets(workbook.Sheets.Count)
    # new_sheet.Name = "new_sheet"

    # 复制excel
    # workbook.Copy()

    workbook.SaveAs(output_file)
    print(f"{output_file}已保存成功！")
    excelApp.Quit()


def test2():
    """word制表符"""
    import datetime
    import os

    from docx import Document
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Inches

    def set_custom_tab_stops():
        doc = Document()
        paragraph = doc.add_paragraph()

        # 清除默认制表位
        paragraph.paragraph_format.tab_stops.clear_all()

        # 添加自定义制表位
        # 在2英寸处添加左对齐制表位
        tab_stop1 = paragraph.paragraph_format.tab_stops.add_tab_stop(
            Inches(1.7),
            WD_TAB_ALIGNMENT.LEFT
        )

        # 在4英寸处添加居中对齐制表位
        tab_stop2 = paragraph.paragraph_format.tab_stops.add_tab_stop(
            Inches(3.5),
            WD_TAB_ALIGNMENT.LEFT
        )

        # 在6英寸处添加右对齐制表位
        tab_stop3 = paragraph.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.0),
            WD_TAB_ALIGNMENT.LEFT
        )

        # 使用制表位
        paragraph.add_run("\t")  # 制表符1
        paragraph.add_run("文件编号：")
        paragraph.add_run("\t")  # 制表符2
        paragraph.add_run("C2GM-013-000")
        # paragraph.add_run("\t")  # 跳到4英寸处
        # paragraph.add_run("居中文字")
        # paragraph.add_run("\t")  # 跳到6英寸处
        # paragraph.add_run("右对齐文字")

        return doc

    def set_custom_2():
        from docx import Document
        from docx.shared import Cm
        from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER

        document = Document()
        paragraph = document.add_paragraph('\t制表符')
        paragraph_format = paragraph.paragraph_format
        tab_stops = paragraph_format.tab_stops
        # tab_stop = tab_stops.add_tab_stop(Cm(5.0))  # 插入制表符
        # print(tab_stop.position)  # 1800225
        # print(tab_stop.position.cm)  # 5.000625
        tab_stop = tab_stops.add_tab_stop(Cm(10.0), alignment=WD_TAB_ALIGNMENT.RIGHT,
                                          leader=WD_TAB_LEADER.DOTS)  # 右对齐，前导字符为点
        # tab_stop = tab_stops.add_tab_stop(Cm(10.0), alignment=WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_LEADER.DOTS)  # 右对齐，前导字符为点

        document.save('test.docx')  # 保存

    if __name__ == '__main__':
        current_time = datetime.datetime.now().strftime('%y%m%d%H%M%S')
        file_path = r"D:\Code\Project\tools\data\temp"
        file = os.path.join(file_path, f'test_{current_time}.docx')

        doc = Document()
        doc.settings.odd_and_even_pages_header_footer = True

        for section in doc.sections:
            section.different_first_page_header_footer = True
            first_footer = section.first_page_footer
            first_footer.paragraphs[0].add_run("这是首页！")
            # first_footer.paragraphs[0].font.bold = True
            # first_footer.paragraphs[0].add_run().font.size = Pt(36.0)
            first_footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.save(file)


def test3():
    """word插入图像"""
    import io

    from docx import Document
    from PIL import Image
    import os

    from docx.shared import Cm

    input_file = r"D:\Code\Project\tools\data\test\test.docx"
    doc = Document(input_file)

    def extract_pics(block):

        for idx, shape in enumerate(block.paragraphs):
            rid = shape._inline

    # 创建一个新的Word文档
    new_doc = Document()

    # 获取图像目录中的所有文件名
    image_dir = r"D:\Code\Project\tools\data\images"
    images = [f for f in os.listdir(image_dir) if os.path.isfile(os.path.join(image_dir, f))]

    # 将每个图像插入到文档中
    row_num = 0
    col_num = 0
    table = doc.add_table(rows=2, cols=2)
    cell = table.rows[0].cells[0]
    for i, image_name in enumerate(images):
        # 打开图像
        img = Image.open(os.path.join(image_dir, image_name))

        # 将图像转换为BytesIO对象
        byte_arr = io.BytesIO()
        img.save(byte_arr, format='PNG')
        byte_arr.seek(0)

        # 插入图像
        run = cell.add_paragraph().add_run()
        inline_shape = run.add_picture(byte_arr, width=Cm(2), height=Cm(2))

        col_num += 1

        # 如果当前行已经满了（即插入了4张图片），则开始新的一行
        if col_num == 4:
            row_num += 1
            col_num = 0

    # 保存文档
    new_doc.save('output.docx')


def test4():
    """word表格宽高"""
    import datetime
    from docx import Document

    input_file = r"D:\Code\Project\tools\data\test\test2.docx"
    doc = Document(input_file)

    for table in doc.tables:
        for column in table.columns:
            for cell in column.cells:
                print(cell.text)

    for table in doc.tables:
        for idx, column in enumerate(table.columns):
            width = table.columns[idx].width
            print(width)

    def get_cell_size(self, cell, table, row_idx, col_idx):
        """获取单元格的宽高"""
        if cell.width:
            width = cell.width.inches
        else:
            width = table.columns[col_idx].width.inches

        return width


def test5():
    # 抽取 docx 中的图片
    from typing import List
    from docx import Document
    from docx.enum.shape import WD_INLINE_SHAPE_TYPE
    from docx.image.image import Image
    from docx.oxml import CT_Inline, CT_GraphicalObject, CT_GraphicalObjectData, CT_Picture, CT_BlipFillProperties, \
        CT_Blip, \
        CT_R

    def extract_pics(file_path: str) -> None:
        doc = Document(file_path)
        for idx, shape in enumerate(doc.inline_shapes):
            # 获取图片的rid
            rid = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
            # 图片的名称
            name = shape._inline.docPr.name
            # 根据rid获取图片对象
            image_obj = doc.part.rels.get(rid).target_part.image
            print("image:", image_obj)
            print("file_name:", image_obj.filename)
            print("扩展名:", image_obj.ext)
            # print("图片二进制数据:", image_obj.blob)
            print("图片二进制数据hash:", image_obj.sha1)
            print("内容类型:", image_obj.content_type)

            print("像素宽高{}x{}:".format(image_obj.px_width, image_obj.px_height))
            print("分辨率{} x {}".format(image_obj.horz_dpi, image_obj.vert_dpi))
            # 解析图片信息
            width = shape.width,  # Lendth对象  可以继续调用.inches/.pt/.cm
            height = shape.height

            print("\n")

    def extract_shapes(file_path: str) -> None:
        """ 抽取图形 """
        doc = Document(file_path)
        # 在段落中抽取图形
        for idx, para in enumerate(doc.element.body):
            for run in para:
                # 找到run
                if isinstance(run, CT_R):
                    for inner_run in run:
                        tag_name = inner_run.tag.split("}")[1]
                        if tag_name == "AlternateContent":  # lxml.etree._Element对象
                            fallback = inner_run[1]
                            pict = fallback[0]
                            shape = pict[0]  # 内部图形
                            print("图形数据:", shape.items())
                            # textbox = list(shape) # shape内部是否有文本，需判断
                            # if textbox:
                            #     textbox = textbox[0]
                            #     # 图形内的段落文本
                            #     para_list = list(textbox[0])
                            #     print("图形内部文本:", [para.text for para in para_list])

    if __name__ == '__main__':
        word_path = r"/data/test/13.docx"
        extract_pics(word_path)
        extract_shapes(word_path)


def test6():
    """excel形状"""
    import datetime
    import os
    import time
    import win32com.client as win32
    from PIL import ImageGrab, ImageDraw
    import pythoncom
    import tempfile
    import logging

    from modules.cm_sop_translate.config.config import LOG_PATH
    from modules.common.log import setup_logger

    # 设置日志
    # logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
    # logger = logging.getLogger(__name__)

    logger = setup_logger(log_dir=LOG_PATH, name='logs', level=logging.INFO)

    def combine_excel_shapes_and_images(excel_path, sheet_name, output_folder, overlap_threshold=0.3):
        """
        将Excel工作表中重叠的图片和形状组合并导出为图像

        参数:
        excel_path: Excel文件路径
        sheet_name: 工作表名称
        output_image_path: 输出图像文件路径
        overlap_threshold: 重叠阈值(0-1)，默认30%重叠视为需要组合
        """
        excel = None
        workbook = None
        # 确保输出文件夹存在
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)
        try:
            # 初始化COM (单线程公寓模式)
            pythoncom.CoInitialize()

            # 启动Excel应用程序
            excel = win32.Dispatch("Excel.Application")
            excel.Visible = False  # 不可见模式运行
            excel.DisplayAlerts = False  # 不显示警告
            excel.ScreenUpdating = False  # 禁用屏幕更新以提高性能

            logger.info(f"正在打开Excel文件: {excel_path}")

            # 打开工作簿
            workbook = excel.Workbooks.Open(os.path.abspath(excel_path))
            worksheet = workbook.Worksheets(sheet_name)

            # 激活工作表
            worksheet.Activate()

            # 获取所有形状
            shapes = worksheet.Shapes

            # 如果没有形状，直接返回
            if shapes.Count == 0:
                logger.warning("工作表中没有找到形状或图片")
                return

            logger.info(f"找到 {shapes.Count} 个形状/图片")

            # 收集所有形状的信息
            shape_info = {}
            count = 1
            for i in range(1, shapes.Count + 1):
                try:
                    shape = shapes.Item(i)

                    text = ''
                    try:
                        if hasattr(shape, 'TextEffect'):
                            text = shape.TextEffect.Text
                    except Exception as e:
                        text = str(e)

                    shape_info[i] = {
                        'name': shape.Name,
                        'type': shape.Type,
                        'left': shape.Left,
                        'top': shape.Top,
                        'width': shape.Width,
                        'height': shape.Height,
                        'shape': shape,
                        'text': text
                    }
                    print(f"形状 {i}: {shape.Name} (类型: {shape.Type})")

                    # 复制形状
                    shape.Copy()

                    # 等待剪贴板数据到来
                    time.sleep(0.5)

                    # 从剪贴板获取图像
                    image = ImageGrab.grabclipboard()

                    if image is not None:
                        CT = datetime.datetime.now().strftime('%y%m%d%H%M%S')
                        output_pic = os.path.join(output_folder, f"{shape.Name}_{CT}_{count}.png")
                        image.save(output_pic, 'PNG')
                        print(f"图片保存成功{output_pic}")
                    else:
                        print("未能从剪贴板获取图像")
                except Exception as e:
                    logger.warning(f"无法获取形状 {i} 的信息: {e}")
                    continue

            # 找出重叠的形状
            # overlapping_groups = find_overlapping_shapes(shape_info, overlap_threshold)
            #
            # # 组合重叠的形状
            # combination_count = 1
            # for group in overlapping_groups:
            #     # 生成图片文件名
            #     img_filename = f"image_{combination_count}.png"
            #     img_path = os.path.join(output_folder, img_filename)
            #     if len(group) > 1:
            #         try:
            #
            #             # 找出图片
            #             shape_range = []
            #             for shape_i in group:
            #                 name = shape_info[shape_i]['name']
            #                 if not has_common_char(name, '图片'):
            #                     shape_range.append(shape_info[shape_i])
            #                 else:
            #                     image_shape = shape_info[shape_i]['shape']
            #                     image_bbox = [
            #                         shape_info[shape_i]['left'],
            #                         shape_info[shape_i]['top'],
            #                         shape_info[shape_i]['left'] + shape_info[shape_i]['width'],
            #                         shape_info[shape_i]['top'] + shape_info[shape_i]['height']
            #                     ]
            #
            #             if (len(group) - len(shape_range)) == 1:
            #                 # shape_image=shape_info[shape_i]['shape']
            #                 # 复制形状
            #                 image_shape.Copy()
            #                 # 等待剪贴板数据到来
            #                 time.sleep(0.1)
            #                 # 从剪贴板获取图像
            #                 image = ImageGrab.grabclipboard()
            #
            #                 # 创建一个Draw对象
            #                 draw = ImageDraw.Draw(image)
            #                 for ran in shape_range:
            #                     name1 = ran['name']
            #                     if has_common_char(name, '椭圆'):
            #                         # 绘制椭圆
            #                         ellipse_props = get_ellipse_properties(ran['shape'])
            #                         draw = draw_ellipse_on_image(draw, image_bbox, ellipse_props)
            #                         image.save(img_path, 'PNG')
            #                     elif has_common_char(name, '矩形'):
            #                         # 绘制矩形
            #                         rect_props = get_rectangle_properties(ran['shape'])
            #                         draw = draw_rectangle_on_image(draw, image_bbox, rect_props)
            #                         image.save(img_path, 'PNG')
            #
            #         except Exception as e:
            #             logger.error(f"组合形状时出错: {e}")
            #             continue
            #     else:
            #         one_shape_index = shape_info[group[0]]
            #         sh = one_shape_index['shape']
            #         # print(sh,type(sh))
            #         # 复制形状
            #         sh.Copy()
            #         # 等待剪贴板数据到来
            #         time.sleep(0.1)
            #         # 从剪贴板获取图像
            #         image = ImageGrab.grabclipboard()
            #         if image is not None:
            #             # img_path = os.path.join(output_folder, img_filename)
            #             image.save(img_path, 'PNG')
            #             print("图片保存成功")
            #         else:
            #             print("未能从剪贴板获取图像")
            #     combination_count += 1
            #
            # logger.info(f"总共组合了 {combination_count} 组形状")

            logger.info(f"shape_info: {shape_info}")

        except Exception as e:
            logger.error(f"处理过程中发生错误: {e}")
            # 获取更详细的错误信息
            # try:
            #     import win32api
            #     win32api.FormatMessage(e.excepinfo[5])
            # except:
            #     pass
            # raise e
        finally:
            # 确保清理资源
            try:
                if workbook:
                    workbook.Close(SaveChanges=False)
                    logger.info("工作簿已关闭")
            except Exception as e:
                logger.error(f"关闭工作簿时出错: {e}")

            try:
                if excel:
                    excel.Quit()
                    logger.info("Excel应用程序已退出")
            except Exception as e:
                logger.error(f"退出Excel时出错: {e}")

            # 释放COM资源
            pythoncom.CoUninitialize()

    def find_overlapping_shapes(shape_info, overlap_threshold=0.3):
        """
        找出重叠的形状

        参数:
        shape_info: 形状信息字典
        overlap_threshold: 重叠阈值(0-1)，默认30%重叠视为需要组合

        返回:
        重叠形状的分组列表
        """
        groups = []
        processed = set()

        # 遍历所有形状
        for i in shape_info:
            if i in processed:
                continue

            # 创建一个新组
            group = [i]
            processed.add(i)

            # 查找与当前形状重叠的其他形状
            for j in shape_info:
                if j in processed:
                    continue

                # 检查是否与组内的任何形状重叠
                for shape_id in group:
                    if is_overlapping(shape_info[shape_id], shape_info[j], overlap_threshold):
                        group.append(j)
                        processed.add(j)
                        break

            groups.append(group)

        return groups

    def is_overlapping(shape1, shape2, threshold):
        """
        判断两个形状是否重叠

        参数:
        shape1: 第一个形状的信息
        shape2: 第二个形状的信息
        threshold: 重叠阈值

        返回:
        布尔值，表示是否重叠
        """
        try:
            # 计算两个形状的边界
            left1, top1 = shape1['left'], shape1['top']
            right1, bottom1 = left1 + shape1['width'], top1 + shape1['height']

            left2, top2 = shape2['left'], shape2['top']
            right2, bottom2 = left2 + shape2['width'], top2 + shape2['height']

            # 检查是否有重叠
            if (right1 < left2 or left1 > right2 or
                    bottom1 < top2 or top1 > bottom2):
                return False

            # 计算重叠区域
            overlap_left = max(left1, left2)
            overlap_top = max(top1, top2)
            overlap_right = min(right1, right2)
            overlap_bottom = min(bottom1, bottom2)

            # 计算重叠面积
            overlap_area = (overlap_right - overlap_left) * (overlap_bottom - overlap_top)

            # 计算较小形状的面积
            area1 = shape1['width'] * shape1['height']
            area2 = shape2['width'] * shape2['height']
            min_area = min(area1, area2)

            # 检查重叠比例是否超过阈值
            return (overlap_area / min_area) >= threshold
        except Exception as e:
            logger.error(f"计算形状重叠时出错: {e}")
            return False

    def has_common_char(str1, str2):
        """
        检查两个字符串是否有重复字符（是否包含至少一个相同的字符）

        参数:
        str1 (str): 第一个字符串
        str2 (str): 第二个字符串

        返回:
        bool: 如果存在重复字符返回True，否则返回False
        """
        # 将字符串转换为字符集合
        set1 = set(str1)
        set2 = set(str2)

        # 检查集合是否有交集
        return len(set1 & set2) > 1

    def export_worksheet_as_image(worksheet, output_path):
        """
        将整个工作表导出为图片

        参数:
        worksheet: Excel工作表对象
        output_path: 输出图片路径
        """
        temp_workbook = None
        temp_worksheet = None

        try:
            # 复制当前工作表
            worksheet.Copy(Before=worksheet)

            # 获取复制的临时工作表
            temp_worksheet = worksheet.Application.ActiveSheet

            # 调整页面设置以适应内容
            temp_worksheet.PageSetup.Zoom = False
            temp_worksheet.PageSetup.FitToPagesWide = 1
            temp_worksheet.PageSetup.FitToPagesTall = 1

            # 选择所有使用的单元格
            used_range = temp_worksheet.UsedRange
            used_range.CopyPicture(Format=2)  # 2 = xlPicture

            # 创建一个临时工作表来粘贴图片
            temp_workbook = worksheet.Application.Workbooks.Add()
            temp_worksheet2 = temp_workbook.Worksheets(1)
            temp_worksheet2.Paste()

            # 获取粘贴的图片
            pasted_picture = temp_worksheet2.Shapes.Item(1)

            # 将图片复制到剪贴板
            pasted_picture.Copy()

            # 等待剪贴板内容就绪
            time.sleep(1)

            # 从剪贴板获取图像并保存
            image = ImageGrab.grabclipboard()
            if image:
                # 确保输出目录存在
                os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
                image.save(output_path)
                logger.info(f"图像已保存到: {output_path}")
            else:
                raise Exception("无法从剪贴板获取图像")

        except Exception as e:
            logger.error(f"导出工作表为图像时出错: {e}")
            raise e
        finally:
            # 清理临时工作表
            try:
                if temp_worksheet:
                    temp_worksheet.Application.DisplayAlerts = False
                    temp_worksheet.Delete()
            except:
                pass

            # 清理临时工作簿
            try:
                if temp_workbook:
                    temp_workbook.Close(SaveChanges=False)
            except:
                pass

    def ensure_excel_closed():
        """
        确保所有Excel实例都已关闭
        """
        try:
            # 尝试通过COM获取Excel应用程序
            excel = win32.Dispatch("Excel.Application")
            # 关闭所有工作簿
            for workbook in excel.Workbooks:
                workbook.Close(SaveChanges=False)
            # 退出Excel
            excel.Quit()
            logger.info("已强制关闭所有Excel实例")
        except:
            # 如果无法通过COM关闭，尝试通过任务管理器关闭
            try:
                os.system('taskkill /f /im excel.exe')
                logger.warning("已通过强制方式关闭Excel进程")
            except:
                logger.error("无法关闭Excel进程")

    # 获取 WPS 椭圆 Shape 的属性
    def get_ellipse_properties(shape):
        properties = {
            'left': shape.Left,  # 左侧位置
            'top': shape.Top,  # 顶部位置
            'width': shape.Width,  # 宽度
            'height': shape.Height,  # 高度
            'fill_color': shape.Fill.ForeColor.RGB if shape.Fill.Visible else None,  # 填充颜色
            'line_color': shape.Line.ForeColor.RGB if shape.Line.Visible else None,  # 线条颜色
            'line_weight': shape.Line.Weight if shape.Line.Visible else 0  # 线条粗细
        }
        return properties

    # 获取矩形 Shape 的属性
    def get_rectangle_properties(shape):
        properties = {
            'left': shape.Left,  # 左侧位置
            'top': shape.Top,  # 顶部位置
            'width': shape.Width,  # 宽度
            'height': shape.Height,  # 高度
            'rotation': shape.Rotation,  # 旋转角度
            'fill_color': shape.Fill.ForeColor.RGB if shape.Fill.Visible else None,  # 填充颜色
            'line_color': shape.Line.ForeColor.RGB if shape.Line.Visible else None,  # 线条颜色
            'line_weight': shape.Line.Weight if shape.Line.Visible else 1,  # 线条粗细
            'line_style': shape.Line.Style if hasattr(shape.Line, 'Style') else 1,  # 线条样式
            'rounded_corners': shape.RoundedCorners if hasattr(shape, 'RoundedCorners') else False,  # 圆角
            'corner_radius': shape.CornerRadius if hasattr(shape, 'CornerRadius') else 0,  # 圆角半径
        }
        return properties

    # 在目标图片上绘制椭圆
    def draw_ellipse_on_image(draw, image_bbox, ellipse_props):
        # 打开目标图片
        # draw = ImageDraw.Draw(image)

        # 转换颜色格式（从BGR到RGB）
        fill_color = None
        if ellipse_props['fill_color']:
            bgr = ellipse_props['fill_color']
            fill_color = (bgr & 0xFF, (bgr >> 8) & 0xFF, (bgr >> 16) & 0xFF)

        line_color = 'black'  # 默认黑色
        if ellipse_props['line_color']:
            bgr = ellipse_props['line_color']
            line_color = (bgr & 0xFF, (bgr >> 8) & 0xFF, (bgr >> 16) & 0xFF)

        # 绘制椭圆
        bbox = [
            int(ellipse_props['left'] - image_bbox[0]) + 10,
            int(ellipse_props['top'] - image_bbox[1]) + 5,
            int(ellipse_props['left'] - image_bbox[0] + ellipse_props['width']) + 15,
            int(ellipse_props['top'] - image_bbox[1] + ellipse_props['height']) + 15
        ]

        if fill_color:
            draw.ellipse(bbox, fill=fill_color, outline=line_color, width=int(ellipse_props['line_weight']))
        else:
            draw.ellipse(bbox, outline=line_color, width=max(int(ellipse_props['line_weight']), 1))

        # image.save('combined_shape2.png', 'PNG')
        return draw

    # 简化版本 - 只绘制基本矩形
    def draw_rectangle_on_image(draw, image_bbox, rect_props):
        # draw = ImageDraw.Draw(img)
        # 转换颜色
        fill_color = None
        if rect_props.get('fill_color'):
            bgr = rect_props['fill_color']
            fill_color = (bgr & 0xFF, (bgr >> 8) & 0xFF, (bgr >> 16) & 0xFF)

        line_color = 'black'
        if rect_props.get('line_color'):
            bgr = rect_props['line_color']
            line_color = (bgr & 0xFF, (bgr >> 8) & 0xFF, (bgr >> 16) & 0xFF)

        # 绘制矩形
        left = int(rect_props['left'] - image_bbox[0]) + 10
        top = int(rect_props['top'] - image_bbox[1]) + 5
        right = int(left + rect_props['width']) + 15
        bottom = int(top + rect_props['height']) + 15

        draw.rectangle(
            [left, top, right, bottom],
            fill=fill_color,
            outline=line_color,
            width=int(rect_props.get('line_weight', 1))
        )

        return draw

    # 使用示例
    if __name__ == "__main__":
        # 设置文件路径
        input_file = r"D:\Code\Project\tools\data\test\3.xlsx"
        sheet_name = "M0901"
        output_folder = r"D:\Code\Project\tools\data\temp\pic2"

        # 检查文件是否存在
        if not os.path.exists(input_file):
            logger.error(f"Excel文件 '{input_file}' 不存在")
        else:
            # 确保没有残留的Excel进程
            ensure_excel_closed()

            # 等待一段时间确保Excel完全关闭
            time.sleep(2)

            # 调用函数处理Excel文件
            try:
                combine_excel_shapes_and_images(input_file, sheet_name, output_folder)
            except Exception as e:
                logger.error(f"处理失败: {e}")
                # 确保Excel进程被关闭
                ensure_excel_closed()


def test7():
    """excel形状"""
    import datetime
    import os

    import win32com.client as win32
    from PIL import ImageGrab
    import time

    from openpyxl.reader.excel import load_workbook

    # 启动Excel应用
    excel = win32.gencache.EnsureDispatch('Excel.Application')
    excel.Visible = True  # 设置为True以便观察，False则后台运行

    input_file = r"D:\Code\Project\tools\data\test\3.xlsx"
    output_folder = r"D:\Code\Project\tools\data\temp\pic"

    # 打开工作簿
    workbook = excel.Workbooks.Open(input_file)
    # sheet = workbook.ActiveSheet

    sheet = workbook.Worksheets("M0901")
    # sheet = workbook.Worksheets("Sheet1")

    # 假设组合形状的名称是“Group 1”，或者通过索引获取
    # 如果知道名称，可以直接使用sheet.Shapes("Group 1")
    # 这里以索引0为例，即第一个形状
    # TextBox 7
    # Group 8
    count = 1

    for shape in sheet.Shapes:
        # shape = sheet.Shapes(1)  # 注意：形状索引从1开始
        text = ''
        try:
            if hasattr(shape, 'TextEffect'):
                text = shape.TextEffect.Text
        except Exception as e:
            text = str(e)
        print(f"1***{shape} 2***{shape.Name} 3***{text} 4***{type(shape)}")
        # 复制形状
        shape.Copy()

        # 等待剪贴板数据到来
        time.sleep(0.5)

        # 从剪贴板获取图像
        image = ImageGrab.grabclipboard()

        if image is not None:
            CT = datetime.datetime.now().strftime('%y%m%d%H%M%S')
            output_pic = os.path.join(output_folder, f"{shape.Name}_{CT}_{count}.png")
            image.save(output_pic, 'PNG')
            print(f"图片保存成功{output_pic}")
        else:
            print("未能从剪贴板获取图像")
        count += 1

    # 关闭工作簿和应用
    workbook.Close(SaveChanges=False)
    excel.Quit()


