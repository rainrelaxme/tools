#!/usr/bin/env python3
# -*- coding:utf-8 -*-
"""
@Project: tools.py
@File   : cm_sop_translate.py
@Version:
@Author : RainRelaxMe
@Date   : 2025/9/24 02:00
@Info   : 实现word文档的翻译，包括段落、表格，并将翻译后的文本置于原文本后，且保持原文档格式。其中doc转换为doc利用了win32包，仅支持在windows系统。
"""

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import RGBColor, Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
import datetime
import os
import getpass
from win32com import client as wc

from src.view.professional_project.sop_translate.template import apply_cover_template, apply_preamble_format, apply_approveTable_format
from src.view.professional_project.sop_translate.translate_by_deepseek import Translator
from config.config import VALID_ACCOUNTS


class DocContent:
    def __init__(self):
        self.title = ''

    def get_content(self, input_path):
        """
        读取word文件的内容，返回位置、内容、格式，记录在dataform中
        保持段落和表格在文档中的原始顺序
        """
        doc = Document(input_path)
        data = []

        # 先处理所有段落，记录它们在文档中的位置
        para_positions = {}
        for para_index, para in enumerate(doc.paragraphs):
            # 获取段落在XML中的位置
            para_xml = para._element
            parent = para_xml.getparent()
            if parent is not None:
                position = list(parent).index(para_xml)
                para_positions[(position, para.text)] = para_index

        # 处理所有表格，记录它们在文档中的位置
        table_positions = {}
        for table_index, table in enumerate(doc.tables):
            table_xml = table._element
            parent = table_xml.getparent()
            if parent is not None:
                position = list(parent).index(table_xml)
                table_positions[position] = table_index

        # 按照文档顺序处理所有元素
        all_elements = []

        # 收集所有段落位置
        for (pos, text), para_index in para_positions.items():
            all_elements.append(('paragraph', pos, para_index, text))

        # 收集所有表格位置
        for pos, table_index in table_positions.items():
            all_elements.append(('table', pos, table_index, None))

        # 按位置排序
        all_elements.sort(key=lambda x: x[1])

        # 按排序后的顺序处理元素
        for element_type, position, index, text in all_elements:
            if element_type == 'paragraph':
                para = doc.paragraphs[index]
                para_data = {
                    'type': 'paragraph',
                    'index': position,
                    'element_index': index,
                    'section': 'body',
                    'flag': '',
                    'text': para.text,
                    'para_format': self.get_paragraph_format(para),
                    'runs': self.get_run_format(para)
                }
                data.append(para_data)

            elif element_type == 'table':
                table = doc.tables[index]
                table_data = {
                    'type': 'table',
                    'index': position,
                    'element_index': index,
                    'section': 'body',
                    'flag': '',
                    'table_alignment': WD_TABLE_ALIGNMENT.CENTER,  # 表格居中，非内容居中
                    'rows': self.get_table_content(table),
                    'cols': len(table.columns),
                }
                data.append(table_data)

        return data

    def get_cover_content(self, content_data):
        """获取word中封面的内容：修订记录表格之前内容算作封面"""
        # 先判断标题是表格还是文本，如果是表格那么首页就截止到第二个表格
        cover_content = []
        top_title_type = ''
        cover_table_num = 1
        for index, item in enumerate(content_data):
            if item['flag'] == 'top_title':
                top_title_type = item['type']
            if top_title_type == 'table':
                cover_table_num = 2
            if item['type'] == 'table' and item['element_index'] == cover_table_num:
                break
            # 先判断再记录
            cover_content.append(item)

        return cover_content

    def get_paragraph_format(self, paragraph):
        """提取段落格式信息"""
        return {
            'style': paragraph.style.name if paragraph.style else 'Normal',
            'alignment': str(paragraph.alignment) if paragraph.alignment else None,
            'space_before': paragraph.paragraph_format.space_before,
            'space_after': paragraph.paragraph_format.space_after,
            'line_spacing': paragraph.paragraph_format.line_spacing,
            'first_line_indent': paragraph.paragraph_format.first_line_indent,
            'left_indent': paragraph.paragraph_format.left_indent
        }

    def get_run_format(self, paragraph):
        # 获取段落格式
        runs_data = []
        for run in paragraph.runs:
            run_format = {
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_size': run.font.size.pt if run.font.size else None,
                'font_name': run.font.name,
                'font_color': run.font.color.rgb if run.font.color and run.font.color.rgb else None,
                'font_color_theme': run.font.color.theme_color if run.font.color else None
            }

            runs_data.append(run_format)

        return runs_data

    def get_table_content(self, table):
        """
        获取表格内容
        """
        rows = []
        cells = []
        for row_index, row in enumerate(table.rows):
            row_data = {'cells': []}
            for cell_index, cell in enumerate(row.cells):

                # 只能判断横向合并,还有种方法，通过合并的两个单元格相等判断
                cells.append(cell)
                grid_span = int(cell.grid_span) if cell.grid_span else 1  # 单元格的跨度，若是合并单元格则大于1
                is_merge_start = False
                if grid_span > 1:
                    is_merge_start = True
                    if len(cells) > 1 and cell == cells[-2]:  # 判断是否等于前一个单元格，即是否与其是合并单元格。
                        is_merge_start = False

                cell_content = {
                    'row': row_index,
                    'col': cell_index,
                    'width': cell.width.inches,
                    'grid_span': grid_span,
                    'is_merge_start': is_merge_start,
                    'text': cell.text,
                    'paragraphs': []
                }

                if cell.text.strip():
                    for cell_para_index, cell_para in enumerate(cell.paragraphs):
                        cell_para_format = self.get_paragraph_format(cell_para)
                        para_data = {
                            'para_index': cell_para_index,
                            'text': cell_para.text,
                            'para_format': cell_para_format,
                            'runs': self.get_run_format(cell_para)
                        }
                        cell_content['paragraphs'].append(para_data)

                row_data['cells'].append(cell_content)
            rows.append(row_data)

        return rows

    def flag_title(self, content_data):
        """
        标记文件大标题：第一个非空内容，表格、段落都可以。
        """
        # 若是标题是用表格做的，还要处理
        title_text = ''
        for index, item in enumerate(content_data):
            if item['type'] == 'paragraph' and item['text'].strip():
                content_data[index]['flag'] = 'top_title'
                title_text = item['text'].strip()
                break
            if item['type'] == 'table' and (len(item['rows']) == 1 and item['cols'] == 1):
                # item['rows'][0]['cells'][0]['text']
                content_data[index]['flag'] = 'top_title'
                title_text = item['rows'][0]['cells'][0]['text'].strip()
                break

    def flag_preamble(self, content_data):
        """
        标记文档的头信息，如文件编号、版本、制定部门、制定日期、总页次。
        """
        # 找到文件头信息的位置，从标题后第一个有内容的段落
        # 先只考虑大标题是段落的形式
        for index, item in enumerate(content_data):
            if item['flag'] == 'top_title':
                title_end_index = index
            if item['type'] == 'table':
                preamble_end_index = index
        cut_title_data = content_data[title_end_index+1:preamble_end_index]

        for index, item in enumerate(cut_title_data):
            if item['type'] == 'paragraph' and item['text'].strip():
                cut_title_data[index]['flag'] = 'preamble'

    def flag_approveTable(self, content_data):
        """
        标记文档封面的审批表格
        """
        for item in content_data:
            if item['type'] == 'table' and item['element_index'] == 0:
                item['flag'] = 'approve'


def set_paper_size(doc):
    """
    设置纸张大小和页边距:A4,页边距2/1.1/2/2
    """
    # 获取第一个节（默认存在）
    section = doc.sections[0]

    # 设置纸张方向（纵向或横向）
    section.orientation = WD_ORIENTATION.PORTRAIT  # PORTRAIT: 纵向, LANDSCAPE: 横向

    # 设置纸张大小（A4）
    section.page_width = Cm(21)  # A4宽度 21cm
    section.page_height = Cm(29.7)  # A4高度 29.7cm

    # 设置页边距（上下左右）
    section.top_margin = Cm(2)  # 上边距
    section.bottom_margin = Cm(1.1)  # 下边距
    section.left_margin = Cm(2)  # 左边距
    section.right_margin = Cm(2)  # 右边距


def create_new_document(content_data, output_path, translations=None):
    """
    根据记录的内容和格式生成新的Word文档
    """
    doc = Document()
    set_paper_size(doc)

    for index, item in enumerate(content_data['content_data']):
        if item['type'] == 'paragraph':
            # 创建新段落
            paragraph = doc.add_paragraph()

            # 应用文件头信息格式
            if item['flag'] == 'preamble':
                apply_preamble_format(paragraph, item)
            else:
                # 应用段落格式
                apply_paragraph_format(paragraph, item['para_format'])

                # 应用运行文本、格式
                for run_data in item['runs']:
                    run = paragraph.add_run(run_data['text'])
                    apply_run_format(run, run_data)

        elif item['type'] == 'table':
            # 创建表格
            # 在表格前添加分页符
            if item['element_index'] > 1:
                page_break_para = doc.add_paragraph()
                run = page_break_para.add_run()
                run.add_break(WD_BREAK.PAGE)

            table = doc.add_table(rows=len(item['rows']), cols=item['cols'])
            table.style = 'Table Grid'

            # 设置表格样式
            # 设置审批表格样式
            if item['flag'] == 'approve':
                apply_approveTable_format(table)

            apply_table_format(table, item)

            # 应用表格内容
            for row_idx, row_data in enumerate(item['rows']):
                for cell_idx, cell_data in enumerate(row_data['cells']):
                    if (cell_data['grid_span'] > 1 and cell_data['is_merge_start']) or cell_data['grid_span'] == 1:
                        cell = table.rows[row_idx].cells[cell_idx]

                        # 清空默认段落
                        for paragraph in cell.paragraphs:
                            p = paragraph._element
                            p.getparent().remove(p)

                        # 添加内容到单元格
                        if cell_data['paragraphs']:
                            for para_data in cell_data['paragraphs']:
                                cell_para = cell.add_paragraph()
                                apply_paragraph_format(cell_para, para_data['para_format'])

                                for run_data in para_data['runs']:
                                    run = cell_para.add_run(run_data['text'])
                                    apply_run_format(run, run_data)
                        else:
                            # 如果没有详细的段落信息，只添加文本
                            cell.text = cell_data['text']

    # 保存文档
    doc.save(output_path)
    print(f"新文档已保存到: {output_path}")


def apply_paragraph_format(paragraph, format_info):
    """应用段落格式"""
    # 使用安全的get方法访问字典，避免KeyError
    if format_info.get('alignment'):
        alignment_map = {
            'CENTER': WD_ALIGN_PARAGRAPH.CENTER,
            'CENTER (1)': WD_ALIGN_PARAGRAPH.CENTER,
            'RIGHT': WD_ALIGN_PARAGRAPH.RIGHT,
            'JUSTIFY': WD_ALIGN_PARAGRAPH.JUSTIFY,
            'LEFT': WD_ALIGN_PARAGRAPH.LEFT
        }
        alignment = format_info['alignment'].split('.')[-1] if '.' in format_info['alignment'] else format_info[
            'alignment']
        paragraph.alignment = alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)

    # 使用get方法安全访问，如果键不存在则返回None
    space_before = format_info.get('space_before')
    space_after = format_info.get('space_after')
    line_spacing = format_info.get('line_spacing')
    first_line_indent = format_info.get('first_line_indent')
    left_indent = format_info.get('left_indent')

    if space_before is not None:
        paragraph.paragraph_format.space_before = space_before

    if space_after is not None:
        paragraph.paragraph_format.space_after = space_after
    else:
        paragraph.paragraph_format.space_after = 0  # 段落后间距None会默认设置为10，需设置为0

    if line_spacing is not None:
        paragraph.paragraph_format.line_spacing = line_spacing
    if first_line_indent is not None:
        paragraph.paragraph_format.first_line_indent = first_line_indent
    if left_indent is not None:
        paragraph.paragraph_format.left_indent = left_indent


def apply_run_format(run, run_data):
    """应用运行文本格式"""
    # 使用get方法安全访问
    run.bold = run_data.get('bold', False)
    run.italic = run_data.get('italic', False)
    run.underline = run_data.get('underline', False)

    font_size = run_data.get('font_size')
    if font_size:
        run.font.size = Pt(font_size)

    font_name = run_data.get('font_name')
    if font_name:
        run.font.name = font_name
        run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    # 无font_name时采用默认字体
    if run.text.strip():
        run.font.name = 'Times New Roman'  # 设置西文字体
        run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 设置中文字体
    else:
        run.font.name = u'宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 设置中文字体

    font_color = run_data.get('font_color')
    if font_color:
        try:
            # 将十六进制颜色转换为RGB
            # hex_color = font_color.lstrip('#')
            # rgb = tuple(int(hex_color[i:i + 2], 16) for i in (0, 2, 4))
            # run.font.color.rgb = RGBColor(*rgb)
            run.font.color.rgb = RGBColor(*font_color)

        except:
            pass


def apply_table_format(table, table_format_info):
    """
    处理表格样式
    """
    # 自动调整列宽：禁用
    table.autofit = False

    # 表格对齐方式，非内容对齐方式
    table.alignment = table_format_info['table_alignment']

    # 单元格样式
    for row in table_format_info['rows']:
        for cell in row['cells']:
            row = cell['row']
            col = cell['col']
            # 设置表格宽度
            if hasattr(cell, 'width'):
                table.cell(row, col).width = Inches(cell['width'])

            # 合并单元格
            if cell['is_merge_start']:
                row = cell['row']
                col = cell['col']
                # 处理水平合并（gridSpan）
                if cell['grid_span'] > 1:
                    # 合并水平单元格
                    start_cell = table.rows[row].cells[col]
                    end_cell = table.rows[row].cells[col + cell['grid_span'] - 1]
                    start_cell.merge(end_cell)
    # if table_format_info['flag'] == 'top_title':
    #     table.autofit = True


def add_paragraph_translation(original_data, translator, language):
    """添加翻译段落"""
    new_data = []

    for item in original_data:
        new_data.append(item)
        if item['type'] == 'paragraph':
            original_text = item['text']
            if original_text.strip():
                # 在data里添加翻译内容
                # 多种语言
                for lang in language:
                    new_item = item.copy()

                    translated_text = translator.translate(original_text, language=lang, display=True)
                    # translated_text = original_text
                    new_item['text'] = translated_text
                    new_item['language'] = lang

                    # 更改run里面内容
                    new_run = item['runs'][0].copy()
                    new_run['text'] = translated_text
                    new_runs = [new_run]
                    new_item['runs'] = new_runs
                    new_data.append(new_item)

    return new_data


def add_table_translation(original_data, translator, language):
    new_data = original_data.copy()
    for item in new_data:
        if item['type'] == 'table':
            for row in item['rows']:
                for cell in row['cells']:
                    new_cell = cell.copy()
                    new_cell['paragraphs'] = []

                    for para in cell['paragraphs']:
                        original_text = para['text']
                        if original_text.strip():
                            if (cell['grid_span'] > 1 and cell['is_merge_start']) or cell['grid_span'] == 1:
                                new_cell['paragraphs'].append(para)

                                # 多种语言
                                for lang in language:
                                    new_para = para.copy()

                                    translated_text = translator.translate(original_text, language=lang, display=True)
                                    # translated_text = original_text
                                    new_para['text'] = translated_text
                                    new_para['language'] = lang

                                    # 更改run里面内容
                                    new_run = para['runs'][0].copy()
                                    new_run['text'] = translated_text
                                    new_runs = [new_run]
                                    new_para['runs'] = new_runs

                                    new_cell['paragraphs'].append(new_para)
                    cell['paragraphs'] = new_cell['paragraphs']

    return new_data


def add_cover_translation(cover_data, translator, language):
    """封面的'文件编号：C2GM-Z13-000'这种需要通过切分分号来分开翻译"""
    # 获取封面内容
    for index, item in enumerate(cover_data):
        if item['flag'] == 'top_title':
            title_end_index = index+1
        if item['flag'] == 'preamble':
            preamble_end_index = index+1

    # 1. '文件编号'之前的正常翻译
    cover_data_1 = cover_data[:title_end_index]
    after_para_1 = add_paragraph_translation(cover_data_1, translator, language)
    after_table_1 = add_table_translation(after_para_1, translator, language)

    # 2. 文件头信息内容：文件编号之后到审批表格前
    cover_data_2 = cover_data[title_end_index:preamble_end_index]
    translated_cover_data_2 = []
    for item in cover_data_2:
        translated_cover_data_2.append(item)
        if item['type'] == 'paragraph' and item['flag'] == 'preamble':
            original_text = item['text']

            # 根据冒号拆分段落
            split_text = []
            if ":" in original_text:
                split_text = original_text.split(":", 1)    # 1代表只分割第一个冒号
            elif "：" in original_text:
                split_text = original_text.split("：", 1)    # 1代表只分割第一个冒号

            if split_text:
                for lang in language:
                    new_item = item.copy()

                    translated_text = ''
                    for part in split_text:
                        part_translated = translator.translate(part, language=lang, display=True)
                        translated_text = translated_text + part_translated + '：'
                    translated_text = translated_text[:-1]
                    # print(translated_text)
                    new_item['text'] = translated_text
                    new_item['language'] = lang

                    # 更改run里面内容
                    new_run = item['runs'][0].copy()
                    new_run['text'] = translated_text
                    new_runs = [new_run]
                    new_item['runs'] = new_runs
                    translated_cover_data_2.append(new_item)

    # 3. 审批表格翻译
    cover_data_3 = cover_data[preamble_end_index:]
    after_para_3 = add_paragraph_translation(cover_data_3, translator, language)
    after_table_3 = add_table_translation(after_para_3, translator, language)

    # 4. 汇总
    new_cover_data = []
    for item in after_table_1:
        new_cover_data.append(item)
    for item in translated_cover_data_2:
        new_cover_data.append(item)
    for item in after_table_3:
        new_cover_data.append(item)

    return new_cover_data


def doc_to_docx(doc_path, docx_path=None):
    """
    将doc文件转换为docx格式
    """
    # 如果未指定输出路径，自动生成
    if docx_path is None:
        docx_path = doc_path.replace('.doc', '.docx')

    # 启动Word应用程序
    word = wc.Dispatch('Word.Application')
    word.Visible = False  # 不显示Word界面

    try:
        # 打开doc文档
        doc = word.Documents.Open(doc_path)
        # 另存为docx格式
        doc.SaveAs(docx_path, FileFormat=16)  # 16表示docx格式
        doc.Close()
        print(f"转换成功: {doc_path} -> {docx_path}")
        return docx_path
    except Exception as e:
        print(f"转换失败: {e}")
        return False
    finally:
        word.Quit()


def login():
    """登录验证函数"""
    max_attempts = 3
    attempts = 0

    while attempts < max_attempts:
        username = input("请输入用户名: ").strip()
        password = getpass.getpass("请输入密码: ").strip()

        # 验证账号密码
        if username in VALID_ACCOUNTS and VALID_ACCOUNTS[username] == password:
            print(f"\n✅ 登录成功！欢迎 {username}！")
            return True
        else:
            attempts += 1
            remaining_attempts = max_attempts - attempts
            print(f"\n❌ 用户名或密码错误！剩余尝试次数: {remaining_attempts}")

            if remaining_attempts > 0:
                print("请重新输入...")
            else:
                print("\n🚫 登录失败次数过多，程序退出！")
                return False

    return False


def check_license():
    """简单的许可证检查（可选功能）"""
    import datetime as dt
    expiry_date = dt.datetime(2025, 10, 31)  # 设置过期时间

    if dt.datetime.now() > expiry_date:
        print("🚫 软件许可证已过期，请联系管理员！")
        return False
    return True


if __name__ == "__main__":
    # 获取当前时间戳
    current_time = datetime.datetime.now().strftime('%y%m%d%H%M%S')
    language = ['英语', '越南语']

    # input_file = r"D:\Code\Project\tools\data\test2.docx"
    # input_file = r"D:\Code\Project\tools\data\1.C2LG-001-000-A08 供应商管理程序.docx"
    input_file = r"/data/13.C2GM-Z13-000-A00 管理评审程序.docx"
    # input_file = r"F:\Code\Project\tools\data\13.C2GM-Z13-000-A00 管理评审程序.docx"
    # input_file = r"F:\Code\Project\tools\data\1.C2LG-001-000-A08 供应商管理程序.docx"

    output_folder = r"D:\Code\Project\tools\data\temp"
    # output_folder = r"F:\Code\Project\tools\data\temp"

    file_base_name = os.path.basename(input_file)
    output_file = output_folder + "/" + file_base_name.replace(".docx", f"_translate_{current_time}.docx")

    translator = Translator()
    print(f"********************start at {current_time}********************")
    # 1. 读取原文档内容
    new_doc = DocContent()
    content_data = new_doc.get_content(input_file)

    # 2. 标注关键信息：大标题、头信息，同时修改content_data内容
    new_doc.flag_title(content_data)
    new_doc.flag_preamble(content_data)
    new_doc.flag_approveTable(content_data)

    # 3. 获取封面信息
    cover_data = new_doc.get_cover_content(content_data)

    # 4. 翻译
    # ① 翻译封面
    translated_cover_data = add_cover_translation(cover_data, translator, language)
    # ② 翻译其它内容
    body_data = content_data[len(cover_data):]
    after_para = add_paragraph_translation(body_data, translator, language)
    after_table = add_table_translation(after_para, translator, language)
    # ③ 合并
    translated_content_data = []
    for item in translated_cover_data:
        translated_content_data.append(item)
    for item in after_table:
        translated_content_data.append(item)

    # 5. 处理内容
    formatted_content = apply_cover_template(translated_content_data, translated_cover_data)

    # 6. 创建新文档
    create_new_document(formatted_content, output_file)

    print(f"********************end********************")
