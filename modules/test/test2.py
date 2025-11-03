import datetime
import os

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches


def set_custom_tab_stops():
    doc = Document()
    paragraph = doc.add_paragraph()

    # 清除默认制表位
    paragraph.paragraph_format.tab_stops.clear_all()

    # 添加自定义制表位
    # 在2英寸处添加左对齐制表位
    tab_stop1 = paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(1.7),
        WD_TAB_ALIGNMENT.LEFT
    )

    # 在4英寸处添加居中对齐制表位
    tab_stop2 = paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(3.5),
        WD_TAB_ALIGNMENT.LEFT
    )

    # 在6英寸处添加右对齐制表位
    tab_stop3 = paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.0),
        WD_TAB_ALIGNMENT.LEFT
    )

    # 使用制表位
    paragraph.add_run("\t")  # 制表符1
    paragraph.add_run("文件编号：")
    paragraph.add_run("\t")  # 制表符2
    paragraph.add_run("C2GM-013-000")
    # paragraph.add_run("\t")  # 跳到4英寸处
    # paragraph.add_run("居中文字")
    # paragraph.add_run("\t")  # 跳到6英寸处
    # paragraph.add_run("右对齐文字")

    return doc

def set_custom_2():
    from docx import Document
    from docx.shared import Cm
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER

    document = Document()
    paragraph = document.add_paragraph('\t制表符')
    paragraph_format = paragraph.paragraph_format
    tab_stops = paragraph_format.tab_stops
    # tab_stop = tab_stops.add_tab_stop(Cm(5.0))  # 插入制表符
    # print(tab_stop.position)  # 1800225
    # print(tab_stop.position.cm)  # 5.000625
    tab_stop = tab_stops.add_tab_stop(Cm(10.0), alignment=WD_TAB_ALIGNMENT.RIGHT,
                                      leader=WD_TAB_LEADER.DOTS)  # 右对齐，前导字符为点
    # tab_stop = tab_stops.add_tab_stop(Cm(10.0), alignment=WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_LEADER.DOTS)  # 右对齐，前导字符为点

    document.save('test.docx')  # 保存


if __name__ == '__main__':
    current_time = datetime.datetime.now().strftime('%y%m%d%H%M%S')
    file_path = r"D:\Code\Project\tools\data\temp"
    file = os.path.join(file_path, f'test_{current_time}.docx')

    doc = Document()
    doc.settings.odd_and_even_pages_header_footer = True

    for section in doc.sections:
        section.different_first_page_header_footer = True
        first_footer = section.first_page_footer
        first_footer.paragraphs[0].add_run("这是首页！")
        # first_footer.paragraphs[0].font.bold = True
        # first_footer.paragraphs[0].add_run().font.size = Pt(36.0)
        first_footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.save(file)
