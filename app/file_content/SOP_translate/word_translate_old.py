from docx import Document
import requests
import json
import time
import os
from typing import List, Dict, Any
import re
import logging

# 设置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class DeepSeekTranslator:
    def __init__(self, api_key: str):
        self.api_key = api_key
        self.api_url = "https://api.deepseek.com/v1/chat/completions"
        self.headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        self.timeout = 60  # 增加超时时间
        self.max_retries = 5  # 增加重试次数

    def translate_text(self, text: str) -> str:
        """使用DeepSeek API翻译文本"""
        if not text.strip():
            return text

        # 清理文本
        text = text.strip()

        # 如果文本太长，分段处理
        if len(text) > 1500:
            return self._handle_long_text(text)

        prompt = f"""
请将以下中文文本专业地翻译成英文，保持专业术语的准确性：

{text}

只需返回翻译后的英文文本，不要添加任何额外的解释或注释。
对于质量管理体系相关的术语，请使用标准英文翻译。
"""

        payload = {
            "model": "deepseek-chat",
            "messages": [
                {"role": "system", "content": "你是一个专业的翻译专家，擅长质量管理体系文档的翻译。熟悉ISO9001标准术语。"},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.1,
            "max_tokens": 2000
        }

        for attempt in range(self.max_retries):
            try:
                logger.info(f"尝试翻译 (第{attempt + 1}次): {text[:50]}...")

                response = requests.post(
                    self.api_url,
                    headers=self.headers,
                    json=payload,
                    timeout=self.timeout
                )

                if response.status_code == 200:
                    result = response.json()
                    translated_text = result['choices'][0]['message']['content'].strip()

                    # 清理可能的额外内容
                    translated_text = self._clean_translation(translated_text)
                    logger.info(f"翻译成功: {translated_text[:50]}...")
                    return translated_text
                else:
                    logger.warning(
                        f"API错误 {response.status_code}: {response.text}, 重试 {attempt + 1}/{self.max_retries}")
                    time.sleep(2 * (attempt + 1))  # 指数退避

            except requests.exceptions.Timeout:
                logger.warning(f"请求超时, 重试 {attempt + 1}/{self.max_retries}")
                time.sleep(3 * (attempt + 1))  # 超时后等待更久
            except requests.exceptions.ConnectionError:
                logger.warning(f"连接错误, 重试 {attempt + 1}/{self.max_retries}")
                time.sleep(5 * (attempt + 1))  # 连接错误等待更久
            except Exception as e:
                logger.error(f"翻译出错: {e}, 重试 {attempt + 1}/{self.max_retries}")
                time.sleep(2 * (attempt + 1))

        logger.error(f"所有重试都失败，返回原文: {text[:50]}...")
        return text  # 如果所有重试都失败，返回原文

    def _handle_long_text(self, text: str) -> str:
        """处理长文本，分段翻译"""
        # 按段落分割
        paragraphs = text.split('\n')
        translated_paragraphs = []

        for para in paragraphs:
            if para.strip():
                if len(para) > 1000:
                    # 如果单个段落也很长，按句子分割
                    sentences = re.split(r'(?<=[.!?。！？]) +', para)
                    current_chunk = ""
                    translated_sentences = []

                    for sentence in sentences:
                        if len(current_chunk) + len(sentence) < 800:
                            current_chunk += sentence + " "
                        else:
                            if current_chunk:
                                translated = self.translate_text(current_chunk)
                                translated_sentences.append(translated)
                            current_chunk = sentence + " "

                    if current_chunk:
                        translated = self.translate_text(current_chunk)
                        translated_sentences.append(translated)

                    translated_paragraphs.append(" ".join(translated_sentences))
                else:
                    translated = self.translate_text(para)
                    translated_paragraphs.append(translated)

        return "\n".join(translated_paragraphs)

    def _clean_translation(self, text: str) -> str:
        """清理翻译结果中的额外内容"""
        # 移除可能的引号或特殊标记
        text = re.sub(r'^["\']|["\']$', '', text)
        # 移除翻译说明文本
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            if not any(
                    keyword in line.lower() for keyword in ['translat', '翻译', 'note:', '提示', 'here is', '以下是']):
                cleaned_lines.append(line)
        return '\n'.join(cleaned_lines).strip()


def translate_docx_with_bilingual_content(input_path: str, output_path: str, api_key: str):
    """翻译整个docx文件，生成双语对照版本，保留原有中文内容"""

    # 初始化翻译器
    translator = DeepSeekTranslator(api_key)

    # 加载文档
    doc = Document(input_path)

    # 创建备份
    backup_path = output_path + ".backup.docx"
    doc.save(backup_path)
    logger.info(f"已创建备份文件: {backup_path}")

    # 统计需要翻译的内容
    paragraphs_to_translate = []
    tables_to_translate = []

    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            paragraphs_to_translate.append((i, para))

    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                if cell.text.strip():
                    tables_to_translate.append((table_idx, row_idx, cell_idx, cell))

    total_items = len(paragraphs_to_translate) + len(tables_to_translate)

    logger.info(f"开始翻译文档: {input_path}")
    logger.info(
        f"需要翻译的内容: {len(paragraphs_to_translate)}个段落 + {len(tables_to_translate)}个表格单元格 = {total_items}项")

    # 翻译段落 - 在原文后添加英文翻译
    if paragraphs_to_translate:
        logger.info("开始翻译段落...")

        for idx, (para_idx, para) in enumerate(paragraphs_to_translate, 1):
            try:
                original_text = para.text
                logger.info(f"[段落 {idx}/{len(paragraphs_to_translate)}] {original_text[:30]}...")

                translated_text = translator.translate_text(original_text)

                # 在原文后添加英文翻译，用换行分隔
                para.text = f"{original_text}\n{translated_text}"

                logger.info(f"✓ 双语对照完成: {translated_text[:30]}...")

                # 定期保存进度
                if idx % 10 == 0:
                    doc.save(output_path)
                    logger.info(f"已保存进度 ({idx}/{len(paragraphs_to_translate)})")

                time.sleep(1)  # 避免频率限制

            except Exception as e:
                logger.error(f"段落翻译失败: {e}")
                continue

    # 翻译表格 - 在原文后添加英文翻译
    if tables_to_translate:
        logger.info("开始翻译表格...")

        for idx, (table_idx, row_idx, cell_idx, cell) in enumerate(tables_to_translate, 1):
            try:
                original_text = cell.text
                logger.info(
                    f"[表格 {idx}/{len(tables_to_translate)}] 表{table_idx + 1}-行{row_idx + 1}-列{cell_idx + 1}: {original_text[:20]}...")

                translated_text = translator.translate_text(original_text)

                # 在原文后添加英文翻译，用换行分隔
                cell.text = f"{original_text}\n{translated_text}"

                logger.info(f"✓ 双语对照完成: {translated_text[:20]}...")

                # 定期保存进度
                if idx % 20 == 0:
                    doc.save(output_path)
                    logger.info(f"已保存进度 ({idx}/{len(tables_to_translate)})")

                time.sleep(0.8)  # 表格翻译可以稍快一些

            except Exception as e:
                logger.error(f"表格翻译失败: {e}")
                continue

    # 最终保存
    doc.save(output_path)
    logger.info(f"双语对照翻译完成！文件已保存至: {output_path}")

    # 统计成功翻译的数量
    success_paragraphs = sum(1 for i, para in paragraphs_to_translate if '\n' in para.text)
    success_tables = sum(1 for item in tables_to_translate if '\n' in item[3].text)

    logger.info(
        f"翻译统计: {success_paragraphs}/{len(paragraphs_to_translate)} 段落, {success_tables}/{len(tables_to_translate)} 表格单元格")


def check_api_connectivity(api_key: str) -> bool:
    """检查API连接性"""
    try:
        response = requests.get(
            "https://api.deepseek.com/v1/models",
            headers={"Authorization": f"Bearer {api_key}"},
            timeout=10
        )
        return response.status_code == 200
    except Exception as e:
        logger.error(f"API连接检查失败: {e}")
        return False


# 使用示例
if __name__ == "__main__":
    # 配置你的DeepSeek API密钥
    DEEPSEEK_API_KEY = "sk-a950c7b878a940be93136750a9a47860"

    input_file = "../../file/1.docx"
    output_file = "../../file/QMS-Z000-A00_Quality_Management_Manual_Bilingual.docx"

    # 开始翻译
    try:
        translate_docx_with_bilingual_content(input_file, output_file, DEEPSEEK_API_KEY)
    except KeyboardInterrupt:
        logger.info("用户中断翻译，已保存当前进度")
    except Exception as e:
        logger.error(f"翻译过程发生严重错误: {e}")
