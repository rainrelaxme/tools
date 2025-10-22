#!/usr/bin/env python3
# -*- coding: utf-8 -*- 
"""
@Project : tools
@File    : test-main.py 
@Author  : Shawn
@Date    : 2025/10/11 9:31 
@Info    : Description of this file

"""
import datetime
import os

from docx import Document

from src.view.production.cm_sop_translate.doc_process import DocumentContent, add_cover_translation, \
    add_paragraph_translation, \
    add_table_translation, add_header_translation, add_footer_translation
from src.view.production.cm_sop_translate.main import create_new_document
from src.view.production.cm_sop_translate.template import apply_cover_template, apply_template
from src.view.production.cm_sop_translate.translator import Translator

if __name__ == "__main__":
    # 获取当前时间戳
    current_time = datetime.datetime.now().strftime('%y%m%d%H%M%S')
    language = ['英语', '越南语']

    input_file = r"F:\Code\Project\tools\data\test\33.docx"

    output_folder = r"F:\Code\Project\tools\data\temp"
    file_base_name = os.path.basename(input_file)
    output_file = output_folder + "/" + file_base_name.replace(".docx", f"_translate_{current_time}.docx")

    translator = Translator()
    print(f"********************start at {current_time}********************")
    # 1. 读取原文档内容
    doc = Document(input_file)
    new_doc = DocumentContent(doc)
    # 正文主体
    content_data = new_doc.get_content(doc)
    # 页眉、页脚
    header_data = new_doc.get_header_content(doc)
    footer_data = new_doc.get_footer_content(doc)
    # picture&shape
    pictures = new_doc.extract_pics(doc)
    shapes = new_doc.extract_shapes(doc)

    # 2. 标注关键信息：大标题、头信息，同时修改content_data内容
    after_title = new_doc.flag_title(content_data)
    after_preamble = new_doc.flag_preamble(content_data)
    after_approve = new_doc.flag_approveTable(content_data)
    after_main_text = new_doc.flag_main_text(content_data)

    # 3. 拆分封面和正文
    data = new_doc.split_cover_body_data(after_main_text)

    # 4. 翻译
    # ① 翻译封面
    translated_cover_data = add_cover_translation(data['cover_data'], translator, language)
    # ② 翻译正文内容
    after_para = add_paragraph_translation(data['body_data'], translator, language)
    translated_body_data = add_table_translation(after_para, translator, language)
    # ③ 翻译页眉、页脚
    translated_header_data = add_header_translation(header_data, translator, language)
    translated_footer_data = add_footer_translation(footer_data, translator, language)

    # 5. 处理内容
    formatted_content = apply_template(body_data=translated_body_data, header_data=translated_header_data, footer_data=translated_footer_data, cover_data=translated_cover_data)

    # 6. 创建新文档
    create_new_document(formatted_content, output_file)

    print(f"********************end********************")