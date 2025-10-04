import datetime
import os

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Inches


def set_custom_tab_stops():
    doc = Document()
    paragraph = doc.add_paragraph()

    # 清除默认制表位
    paragraph.paragraph_format.tab_stops.clear_all()

    # 添加自定义制表位
    # 在2英寸处添加左对齐制表位
    tab_stop1 = paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(1.7),
        WD_TAB_ALIGNMENT.LEFT
    )

    # 在4英寸处添加居中对齐制表位
    tab_stop2 = paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(3.5),
        WD_TAB_ALIGNMENT.LEFT
    )

    # 在6英寸处添加右对齐制表位
    tab_stop3 = paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.0),
        WD_TAB_ALIGNMENT.LEFT
    )

    # 使用制表位
    paragraph.add_run("\t")  # 制表符1
    paragraph.add_run("文件编号：")
    paragraph.add_run("\t")  # 制表符2
    paragraph.add_run("C2GM-013-000")
    # paragraph.add_run("\t")  # 跳到4英寸处
    # paragraph.add_run("居中文字")
    # paragraph.add_run("\t")  # 跳到6英寸处
    # paragraph.add_run("右对齐文字")

    return doc


if __name__ == '__main__':
    doc = set_custom_tab_stops()
    output_folder = r"D:\Code\Project\tools\data\temp"
    current_time = datetime.datetime.now().strftime('%y%m%d%H%M%S')

    file_base_name = 'test2.docx'
    output_file = output_folder + "/" + file_base_name.replace(".docx", f"_translate_{current_time}.docx")

    doc.save(output_file)