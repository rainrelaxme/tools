#!/usr/bin/env python3
# -*- coding: utf-8 -*- 
"""
@Project : tools
@File    : template.py 
@Author  : Shawn
@Date    : 2025/9/30 14:15 
@Info    : Description of this file
"""

import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Cm


def apply_cover_template(content_data, cover_data):
    """应用封面模板"""
    new_cover_data = []

    para_data = {
        'type': 'paragraph',
        'index': None,
        'element_index': None,
        'flag': '',
        'text': '',
        'para_format': {
            'style': 'Normal',
            'alignment': None,
            'space_before': None,
            'space_after': None,
            'line_spacing': 2.0,
            'first_line_indent': None,
            'left_indent': None
        },
        'runs': []
    }
    table_data = {
        'type': 'table',
        'index': None,
        'element_index': None,
        'flag': None,
        'table_alignment': WD_TABLE_ALIGNMENT.CENTER,  # 表格居中，非内容居中
        'rows': None,
        'cols': None,
    }

    # 1. 第一行空格
    para_1 = para_data.copy()
    para_1['index'] = 0
    para_1['element_index'] = 0
    new_cover_data.append(para_1)

    # 2. 大标题
    table_1 = table_data.copy()
    table_1['index'] = 1
    table_1['element_index'] = 0
    table_1['flag'] = 'top_title'
    table_1['cols'] = 1
    rows = []
    for item in cover_data:
        if item['flag'] == 'top_title':
            row = {
                'cells': [
                    {
                        'row': 0,
                        'col': 0,
                        'grid_span': 1,
                        'is_merge_start': False,
                        'text': item['text'],
                        'paragraphs': [
                            {
                                'para_index': 0,
                                'text': item['text'],
                                'para_format': item['para_format'],
                                'runs': item['runs'] if 'runs' in item else []
                            }
                        ]
                    }
                ]
            }
            rows.append(row)
    new_rows = [
        {
            'cells': [
                {
                    'row': rows[0]['cells'][0]['row'],
                    'col': rows[0]['cells'][0]['col'],
                    'grid_span': rows[0]['cells'][0]['grid_span'],
                    'is_merge_start': rows[0]['cells'][0]['is_merge_start'],
                    'text': rows[0]['cells'][0]['text'],
                    'paragraphs': []
                }

            ]
        }
    ]
    for item in rows:
        new_rows[0]['cells'][0]['paragraphs'].append(item['cells'][0]['paragraphs'][0])
    table_1['rows'] = new_rows
    new_cover_data.append(table_1)

    # 3. 接空行
    para_2 = para_data.copy()
    para_2['index'] = 2
    para_2['element_index'] = 1
    new_cover_data.append(para_2)

    # 4. 文件头信息：文档编号等
    index = len(new_cover_data)
    para_index = 2
    i = 0
    for item in cover_data:
        if item['flag'] == 'preamble':
            item['index'] = index
            item['element_index'] = para_index
            new_cover_data.append(item)
            index += 1
            para_index += 1
            i += 1
            if i < 3:  # 翻译语言的种类
                continue

            # 加一个空行
            new_blank = para_data.copy()
            new_blank['index'] = index + 1
            new_blank['element_index'] = para_index + 1
            new_cover_data.append(new_blank)
            index += 1
            para_index += 1
            i = 0

    # 5. 审批表格
    for item in cover_data:
        if item['type'] == 'table':
            item['index'] = len(new_cover_data)+1
            item['element_index'] = 1
            for row in item['rows']:
                for cell in row['cells']:
                    for para in cell['paragraphs']:
                        para['para_format']['line_spacing'] = Pt(20)
            new_cover_data.append(item)

    # 6. 更新content_data
    del content_data[:len(cover_data)]
    content_data[:0] = new_cover_data

    res = {
        "content_data": content_data,
        "cover_data": new_cover_data
    }
    return res


def apply_preamble_format(paragraph, preamble_data):
    """
    应用文件头信息格式，通过制表符对齐
    """
    # 清除默认制表位
    paragraph.paragraph_format.tab_stops.clear_all()

    # 固定行距20磅
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY  # 固定行距
    paragraph_format.line_spacing = Pt(20)
    paragraph_format.space_after = 0

    # 添加自定义制表位
    tab_stop1 = paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(1.7),
        WD_TAB_ALIGNMENT.LEFT
    )
    tab_stop2 = paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(3.5),
        WD_TAB_ALIGNMENT.LEFT
    )

    # 使用制表位
    split_text = preamble_data['text'].split('：')
    paragraph.add_run("\t")  # 制表符1
    run1 = paragraph.add_run(split_text[0]+'：')
    run1.font.bold = True
    run1.font.size = Pt(16.0)
    run1.font.name = u'宋体'
    run1.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 设置中文字体

    paragraph.add_run("\t")  # 制表符2
    run2 = paragraph.add_run(split_text[1].strip())
    run2.font.size = Pt(16.0)
    run2.font.name = 'Times New Roman'


def apply_approveTable_format(table):
    """
    应用审批表格格式
    """
    for i in range(len(table.rows)):
        table.rows[i].height = Cm(2.5)


if __name__ == '__main__':
    doc = Document()
    content_data = {
        "content_data" : [{
            'type': 'paragraph',
            'index': 2,
            'element_index': 1,
            'flag': '',
            'text': '',
            'para_format': {
                'style': 'Normal',
                'alignment': None,
                'space_before': None,
                'space_after': None,
                'line_spacing': 2.0,
                'first_line_indent': None,
                'left_indent': None
            },
            'runs': []
        }, {
            'type': 'paragraph',
            'index': 3,
            'element_index': 2,
            'flag': 'preamble',
            'text': '文件编号：C2GM-Z13-000',
            'para_format': {
                'style': 'Normal',
                'alignment': None,
                'space_before': None,
                'space_after': None,
                'line_spacing': 254000,
                'first_line_indent': 1828800,
                'left_indent': None
            },
            'runs': [{
                'text': '文件编号：',
                'bold': True,
                'italic': None,
                'underline': None,
                'font_size': 16.0,
                'font_name': '宋体',
                'font_color': None,
                'font_color_theme': None
            }, {
                'text': 'C2GM-',
                'bold': True,
                'italic': None,
                'underline': None,
                'font_size': 16.0,
                'font_name': '宋体',
                'font_color': None,
                'font_color_theme': None
            }, {
                'text': 'Z',
                'bold': True,
                'italic': None,
                'underline': None,
                'font_size': 16.0,
                'font_name': '宋体',
                'font_color': None,
                'font_color_theme': None
            }, {
                'text': '13-000',
                'bold': True,
                'italic': None,
                'underline': None,
                'font_size': 16.0,
                'font_name': '宋体',
                'font_color': None,
                'font_color_theme': None
            }
            ]
        }, {
            'type': 'paragraph',
            'index': 4,
            'element_index': 3,
            'flag': 'preamble',
            'text': 'Doc. No.：C2GM-Z13-000',
            'para_format': {
                'style': 'Normal',
                'alignment': None,
                'space_before': None,
                'space_after': None,
                'line_spacing': 254000,
                'first_line_indent': 1828800,
                'left_indent': None
            },
            'runs': [{
                'text': 'Doc. No.：C2GM-Z13-000',
                'bold': True,
                'italic': None,
                'underline': None,
                'font_size': 16.0,
                'font_name': '宋体',
                'font_color': None,
                'font_color_theme': None
            }
            ],
            'language': '英语'
        }, {
            'type': 'paragraph',
            'index': 5,
            'element_index': 4,
            'flag': 'preamble',
            'text': 'Văn bản số：C2GM-Z13-000',
            'para_format': {
                'style': 'Normal',
                'alignment': None,
                'space_before': None,
                'space_after': None,
                'line_spacing': 254000,
                'first_line_indent': 1828800,
                'left_indent': None
            },
            'runs': [{
                'text': 'Văn bản số：C2GM-Z13-000',
                'bold': True,
                'italic': None,
                'underline': None,
                'font_size': 16.0,
                'font_name': '宋体',
                'font_color': None,
                'font_color_theme': None
            }
            ],
            'language': '越南语'
        }, {
            'type': 'paragraph',
            'index': 7,
            'element_index': 6,
            'flag': '',
            'text': '',
            'para_format': {
                'style': 'Normal',
                'alignment': None,
                'space_before': None,
                'space_after': None,
                'line_spacing': 2.0,
                'first_line_indent': None,
                'left_indent': None
            },
            'runs': []
        }, {
            'type': 'paragraph',
            'index': 7,
            'element_index': 6,
            'flag': 'preamble',
            'text': '版    本：      A00',
            'para_format': {
                'style': 'Normal',
                'alignment': None,
                'space_before': None,
                'space_after': None,
                'line_spacing': 254000,
                'first_line_indent': 1828800,
                'left_indent': None
            },
            'runs': [{
                'text': '版    本：',
                'bold': True,
                'italic': None,
                'underline': None,
                'font_size': 16.0,
                'font_name': '宋体',
                'font_color': None,
                'font_color_theme': None
            }, {
                'text': '  ',
                'bold': True,
                'italic': None,
                'underline': None,
                'font_size': 16.0,
                'font_name': '宋体',
                'font_color': None,
                'font_color_theme': None
            }, {
                'text': ' ',
                'bold': True,
                'italic': None,
                'underline': None,
                'font_size': 16.0,
                'font_name': '宋体',
                'font_color': None,
                'font_color_theme': None
            }, {
                'text': '  ',
                'bold': True,
                'italic': None,
                'underline': None,
                'font_size': 16.0,
                'font_name': '宋体',
                'font_color': None,
                'font_color_theme': None
            }, {
                'text': ' ',
                'bold': True,
                'italic': None,
                'underline': None,
                'font_size': 16.0,
                'font_name': '宋体',
                'font_color': None,
                'font_color_theme': None
            }, {
                'text': 'A',
                'bold': True,
                'italic': None,
                'underline': None,
                'font_size': 16.0,
                'font_name': '宋体',
                'font_color': None,
                'font_color_theme': None
            }, {
                'text': '0',
                'bold': True,
                'italic': None,
                'underline': None,
                'font_size': 16.0,
                'font_name': '宋体',
                'font_color': None,
                'font_color_theme': None
            }, {
                'text': '0',
                'bold': True,
                'italic': None,
                'underline': None,
                'font_size': 16.0,
                'font_name': '宋体',
                'font_color': None,
                'font_color_theme': None
            }
            ]
        }, {
            'type': 'paragraph',
            'index': 8,
            'element_index': 7,
            'flag': 'preamble',
            'text': 'Version：A00',
            'para_format': {
                'style': 'Normal',
                'alignment': None,
                'space_before': None,
                'space_after': None,
                'line_spacing': 254000,
                'first_line_indent': 1828800,
                'left_indent': None
            },
            'runs': [{
                'text': 'Version：A00',
                'bold': True,
                'italic': None,
                'underline': None,
                'font_size': 16.0,
                'font_name': '宋体',
                'font_color': None,
                'font_color_theme': None
            }
            ],
            'language': '英语'
        }, {
            'type': 'paragraph',
            'index': 9,
            'element_index': 8,
            'flag': 'preamble',
            'text': 'Ấn bản：A00',
            'para_format': {
                'style': 'Normal',
                'alignment': None,
                'space_before': None,
                'space_after': None,
                'line_spacing': 254000,
                'first_line_indent': 1828800,
                'left_indent': None
            },
            'runs': [{
                'text': 'Ấn bản：A00',
                'bold': True,
                'italic': None,
                'underline': None,
                'font_size': 16.0,
                'font_name': '宋体',
                'font_color': None,
                'font_color_theme': None
            }
            ],
            'language': '越南语'
        }]
    }

    for index, item in enumerate(content_data['content_data']):
        if item['type'] == 'paragraph':
            # 创建新段落
            paragraph = doc.add_paragraph()

            # 应用文件头信息格式
            if item['flag'] == 'preamble':
                apply_preamble_format(paragraph, item)

    doc.save("test.docx")