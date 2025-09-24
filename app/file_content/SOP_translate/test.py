import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.oxml.ns import qn
from SOP_translate_cursor import create_new_document as CND

# 1.记录合并单元格的位置
# 2.应用合并单元格的位置
# 3.创建文档

def get_merge_cells_info(table):
    merged_cells = []

    data = []
    for table_index, table in enumerate(doc.tables):
        data.append([])
        for row_index, row in enumerate(table.rows):
            data[-1].append([])
            for col_index, cell in enumerate(row.cells):
                data[-1][-1].append(cell)
        print(data)


def apply_merged_cells(table, merged_cells_info):
    start_cell = table.rows[0].cells[0]
    end_cell = table.rows[0].cells[2]
    start_cell.merge(end_cell)

    start_cell_vertical = table.rows[1].cells[0]
    end_cell_vertical = table.rows[2].cells[1]
    start_cell_vertical.merge(end_cell_vertical)


def get_table_size(table):
    cells = []
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            # print(f"({row_index}, {col_index})")

            cell_info = {
                'row': row_index,
                'col': col_index,
                'width': cell.width,
                'content': cell.text
            }
            cells.append(cell_info)
    print(cells)


def get_paper_size(doc):
    section = doc.sections[0]
    # 获取的尺寸说是twip，但是换算1cm=567twip，不对。
    paper_size = {
        'orientation': section.orientation,
        'page_width': section.page_width,
        'page_height': section.page_height,
        'top_margin': section.top_margin,
        'bottom_margin': section.bottom_margin,
        'left_margin': section.left_margin,
        'right_margin': section.right_margin,
    }
    print("paper_size", paper_size)

    return paper_size


if __name__ == '__main__':
    datetime = datetime.datetime.now().strftime('%y%m%d%H%M%S')
    input_file = f"../../../file/test.docx"
    output_file = f"../../../file/temp/outputxxx.docx"
    doc = Document(input_file)

    get_paper_size(doc)


    # new_doc = Document()
    #
    # table = new_doc.add_table(rows=5, cols=3)
    # table.style = 'Table Grid'
    #
    # table.cell(0,0).width = Inches(3.6)
    # table.cell(0,1).width = Inches(2.5)
    # table.cell(0,1).width = Inches(2.5)

    # table.cell(1,0).merge(table.cell(1,1))
    # table.cell(1, 0).width = Inches(3)

    # new_doc.save(output_file)
    # print(f"新文档已保存到: {output_file}")




def set_ziti(doc):
    para = doc.add_paragraph()

    run = para.add_run('testn\n你好')
    # run1 = para.add_run('你好，鲍老师')
    # run2 = para.add_run('test text content.')

    # run.font.name = 'Times New Roman'

    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')
    run.font.size = Pt(36)

    # run1.font.name = '宋体'
    # run1.font.size = Pt(36)
    #
    # run2.font.name = '宋体'
    # run2.font.size = Pt(12)

    doc.save(output_file)
