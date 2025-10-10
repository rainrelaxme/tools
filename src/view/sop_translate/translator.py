#!/usr/bin/env python3
# -*- coding: utf-8 -*- 
"""
@Project : tools
@File    : translator.py 
@Author  : Shawn
@Date    : 2025/9/25 11:13 
@Info    : Description of this file
"""

import datetime
import os
import sys

from docx import Document

from src.view.sop_translate.sop_translate import login, check_license, DocContent, add_paragraph_translation, add_table_translation, \
    create_new_document, doc_to_docx, add_cover_translation
from src.view.sop_translate.template import apply_cover_template
from src.view.sop_translate.translate_by_deepseek import Translator


def text_translate(language):
    original_text = input("请输入待翻译内容：\n").strip()
    translator = Translator()
    for lang in language:
        translator.translate(original_text, language=lang)
        # print(f"{lang}: {translator.translate(original_text, language=lang)}")


def docx_translate(language):
    # 设置输入文件夹路径
    input_folder = input("请输入待翻译的文件目录：\n")

    # 设置输出文件夹路径（二级文件夹）
    output_folder = os.path.join(input_folder, "translate_output")
    os.makedirs(output_folder, exist_ok=True)

    # 初始化翻译器
    translator = Translator()


    # 遍历文件夹中的所有文件
    for filename in os.listdir(input_folder):
        print(f"正在处理文件: {filename}")

        # 检查文件是否为Word文档（.docx格式）,若是.doc文件则转为.docx
        if filename.endswith('.doc') and not filename.startswith('~$'):  # 忽略临时文件
            doc_path = os.path.join(input_folder, filename)
            filename = os.path.basename(doc_to_docx(doc_path))

        if filename.endswith('.docx') and not filename.startswith('~$'):
            input_file = os.path.join(input_folder, filename)

            # 生成输出文件名（原文件名+时间）
            current_time = datetime.datetime.now().strftime('%y%m%d%H%M%S')
            file_base_name = os.path.splitext(filename)[0]  # 去掉扩展名
            output_filename = f"{file_base_name}_translate_{current_time}.docx"
            output_file = os.path.join(output_folder, output_filename)

            # 1. 读取原文档内容
            doc = Document(input_file)
            new_doc = DocContent()
            content_data = new_doc.get_content(doc)

            # header_content = new_doc.get_header_content(doc)
            # footer_content = new_doc.get_footer_content(doc)

            # 2. 标注关键信息：大标题、头信息，同时修改content_data内容
            new_doc.flag_title(content_data)
            new_doc.flag_preamble(content_data)
            new_doc.flag_approveTable(content_data)

            # 3. 获取封面信息
            cover_data = new_doc.get_cover_content(content_data)

            # 4. 翻译
            # ① 翻译封面
            translated_cover_data = add_cover_translation(cover_data, translator, language)
            # ② 翻译其它内容
            body_data = content_data[len(cover_data):]
            after_para = add_paragraph_translation(body_data, translator, language)
            after_table = add_table_translation(after_para, translator, language)
            # ③ 合并
            translated_content_data = []
            for item in translated_cover_data:
                translated_content_data.append(item)
            for item in after_table:
                translated_content_data.append(item)

            # 5. 处理内容
            formatted_content = apply_cover_template(translated_content_data, translated_cover_data)

            # 6. 创建新文档
            create_new_document(formatted_content, output_file)
            print(f"  已完成翻译: {output_filename}")

    print(f"所有文件处理完成！输出目录: {output_folder}")




if __name__ == '__main__':

    print("\n" + "=" * 100)
    print("\x20" * 45 + f"文档翻译系统")
    print("=" * 100)
    print("注意：仅支持运行于windows系统！！！")
    # 首先进行登录验证
    # if not login():
    #     exit()

    # 可选：进行许可证检查
    if not check_license():
        exit()

    while True:
        try:
            print("请选择使用方式：\n"
                  "1. 文本翻译\n"
                  "2. 文档翻译\n"
                  "3. 退出")
            option = input().strip()
            language = ['英语', '越南语']
            if option == '1':
                print(f"当前目标语言为{language}")
                text_translate(language)
            elif option == '2':
                print(f"当前目标语言为{language}")
                docx_translate(language)
            elif option == '3':
                print("\n感谢使用！程序即将退出...")
                input("按回车键关闭窗口...")
                sys.exit(0)
            else:
                print("输入错误，请重新输入")

            # print("请选择目标语言（可多选，用\",\"分隔）：\n"
            #       "1. 英语\n"
            #       "2. 越南语")
            # lang_input = input().strip()
            # lang_list = lang_input.split(',')
            # language = []
            # for lang in lang_list:
            #     if lang.strip() == '1':
            #         language.append('英语')
            #     if lang.strip() == '2':
            #         language.append('越南语')
            # print(f"您选择的目标语言为{language}")

        except Exception as e:
            print(f"处理过程中出现错误: {e}")
        current_time = datetime.datetime.now().strftime('%y%m%d%H%M%S')
        print(f"\n********************{current_time} end**********************")
