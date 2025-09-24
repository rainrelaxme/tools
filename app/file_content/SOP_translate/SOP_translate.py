#!/usr/bin/env python3
# -*- coding:utf-8 -*-
"""
@Project: tools.py
@File   : SOP_translate.py
@Version:
@Author : RainRelaxMe
@Date   : 2025/9/24 02:00
@Info   : 实现word文档的翻译，包括段落、表格，并将翻译后的文本置于原文本后，且保持原文档格式
"""

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import RGBColor, Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
import datetime
import os
import re
import getpass

from app.file_content.SOP_translate.translate_by_deepseek import Translator


def get_content(input_path):
    """
    读取word文件的内容，返回位置、内容、格式，记录在dataform中
    保持段落和表格在文档中的原始顺序
    """
    doc = Document(input_path)
    data = []
    # global_index = 0

    # 先处理所有段落，记录它们在文档中的位置
    para_positions = {}
    for para_index, para in enumerate(doc.paragraphs):
        # 获取段落在XML中的位置
        para_xml = para._element
        parent = para_xml.getparent()
        if parent is not None:
            position = list(parent).index(para_xml)
            para_positions[(position, para.text)] = para_index

    # 处理所有表格，记录它们在文档中的位置
    table_positions = {}
    for table_index, table in enumerate(doc.tables):
        table_xml = table._element
        parent = table_xml.getparent()
        if parent is not None:
            position = list(parent).index(table_xml)
            table_positions[position] = table_index

    # 按照文档顺序处理所有元素
    all_elements = []

    # 收集所有段落位置
    for (pos, text), para_index in para_positions.items():
        all_elements.append(('paragraph', pos, para_index, text))

    # 收集所有表格位置
    for pos, table_index in table_positions.items():
        all_elements.append(('table', pos, table_index, None))

    # 按位置排序
    all_elements.sort(key=lambda x: x[1])

    # 按排序后的顺序处理元素
    for element_type, position, index, text in all_elements:
        if element_type == 'paragraph':
            para = doc.paragraphs[index]
            # 获取段落格式
            para_format = get_paragraph_format(para)
            runs_data = []
            for run in para.runs:
                run_format = get_run_format(run)
                runs_data.append(run_format)

            data.append({
                'type': 'paragraph',
                'index': position,
                'element_index': index,
                'text': para.text,
                'para_format': para_format,
                'runs': runs_data
            })
            # global_index += 1

        elif element_type == 'table':
            table = doc.tables[index]
            table_data = {
                'type': 'table',
                'index': position,
                'element_index': index,
                'table_alignment': WD_TABLE_ALIGNMENT.CENTER,  # 表格居中，非内容居中
                'rows': get_table_content(table),
                'cols': len(table.columns),
                'merged_cells': get_merged_cells_info(table),
                'table_format': get_table_format(table),
            }
            data.append(table_data)

    return data


def get_paragraph_format(paragraph):
    """提取段落格式信息"""
    return {
        'style': paragraph.style.name if paragraph.style else 'Normal',
        'alignment': str(paragraph.alignment) if paragraph.alignment else None,
        'space_before': paragraph.paragraph_format.space_before,
        'space_after': paragraph.paragraph_format.space_after,
        'line_spacing': paragraph.paragraph_format.line_spacing,
        'first_line_indent': paragraph.paragraph_format.first_line_indent,
        'left_indent': paragraph.paragraph_format.left_indent
    }


def get_run_format(run):
    """提取运行文本格式信息"""
    return {
        'text': run.text,
        'bold': run.bold,
        'italic': run.italic,
        'underline': run.underline,
        'font_size': run.font.size.pt if run.font.size else None,
        'font_name': run.font.name,
        'font_color': get_rgb_color(run.font.color.rgb) if run.font.color and run.font.color.rgb else None,
        'font_color_theme': run.font.color.theme_color if run.font.color else None
    }


def get_table_content(table):
    """
    获取表格内容
    """
    rows = []
    for row_index, row in enumerate(table.rows):
        row_data = {'cells': []}
        for cell_index, cell in enumerate(row.cells):
            if cell.text.strip():
                cell_content = {
                    'text': cell.text,
                    'paragraphs': []
                }

                for cell_para_index, cell_para in enumerate(cell.paragraphs):
                    cell_para_format = get_paragraph_format(cell_para)
                    cell_runs_data = []
                    for run in cell_para.runs:
                        run_format = get_run_format(run)
                        cell_runs_data.append(run_format)

                    cell_content['paragraphs'].append({
                        'para_index': cell_para_index,
                        'text': cell_para.text,
                        'para_format': cell_para_format,
                        'runs': cell_runs_data
                    })

                row_data['cells'].append(cell_content)
            else:
                row_data['cells'].append({
                    'text': '',
                    'paragraphs': []
                })
        rows.append(row_data)

    return rows


def get_merged_cells_info(table):
    """
    获取合并单元格信息
    """
    merged_cells = []

    # 只能判断横向合并
    """
    # 还有种方法，通过合并的两个单元格相等判断
    data = []
    for table_index, table in enumerate(doc.tables):
        data.append([])
        for row_index, row in enumerate(table.rows):
            data[-1].append([])
            for col_index, cell in enumerate(row.cells):
                data[-1][-1].append(cell)
        print(data)
    """
    cells = []
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            cells.append(cell)
            grid_span = int(cell.grid_span) if cell.grid_span else 1  # 单元格的跨度，即如果合并则大于1
            if grid_span > 1:
                is_merge_start = True
                if len(cells) > 1 and cell == cells[-2]:  # 判断是否等于前一个单元格，即是否与其是合并单元格。
                    is_merge_start = False

                merge_info = {
                    'row': row_index,
                    'col': col_index,
                    'grid_span': grid_span,
                    'is_merge_start': is_merge_start,  # 是否是合并单元格的起点
                }
                merged_cells.append(merge_info)

    return merged_cells


def get_table_format(table):
    """
    获取表格的宽度等信息
    """
    cells = []
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            cell_size = {
                'row': row_index,
                'col': col_index,
                'width': cell.width.inches,
            }
            cells.append(cell_size)

    return cells


def get_rgb_color(rgb_value):
    """将RGB颜色值转换为十六进制字符串"""
    if rgb_value:
        return f"#{rgb_value:06x}"
    return None


def set_paper_size(doc):
    """
    设置纸张大小和页边距:A4,页边距2/1.1/2/2
    """
    # 获取第一个节（默认存在）
    section = doc.sections[0]

    # 设置纸张方向（纵向或横向）
    section.orientation = WD_ORIENTATION.PORTRAIT  # PORTRAIT: 纵向, LANDSCAPE: 横向

    # 设置纸张大小（A4）
    section.page_width = Cm(21)  # A4宽度 21cm
    section.page_height = Cm(29.7)  # A4高度 29.7cm

    # 设置页边距（上下左右）
    section.top_margin = Cm(2)  # 上边距
    section.bottom_margin = Cm(1.1)  # 下边距
    section.left_margin = Cm(2)  # 左边距
    section.right_margin = Cm(2)  # 右边距


def create_new_document(content_data, output_path, translations=None):
    """
    根据记录的内容和格式生成新的Word文档
    """
    doc = Document()
    set_paper_size(doc)

    for item in content_data:
        if item['type'] == 'paragraph':
            # 创建新段落
            paragraph = doc.add_paragraph()

            # 应用段落格式
            apply_paragraph_format(paragraph, item['para_format'])

            # 添加运行文本
            for run_data in item['runs']:
                run = paragraph.add_run(run_data['text'])
                apply_run_format(run, run_data)

        elif item['type'] == 'table':
            # 创建表格
            # 在表格前添加分页符
            if item['element_index'] > 0:
                page_break_para = doc.add_paragraph()
                run = page_break_para.add_run()
                run.add_break(WD_BREAK.PAGE)

            table = doc.add_table(rows=len(item['rows']), cols=item['cols'])
            table.style = 'Table Grid'
            table.alignment = item['table_alignment']

            # 应用合并单元格
            apply_merged_cells(table, item['merged_cells'])
            # 设置表格宽度
            apply_table_format(table, item['table_format'])

            # 填充表格内容
            for row_idx, row_data in enumerate(item['rows']):
                for cell_idx, cell_data in enumerate(row_data['cells']):
                    if cell_idx < len(table.rows[row_idx].cells):
                        cell = table.rows[row_idx].cells[cell_idx]

                        # 清空默认段落
                        for paragraph in cell.paragraphs:
                            p = paragraph._element
                            p.getparent().remove(p)

                        # 添加内容到单元格
                        if cell_data['paragraphs']:
                            for para_data in cell_data['paragraphs']:
                                cell_para = cell.add_paragraph()
                                apply_paragraph_format(cell_para, para_data['para_format'])

                                for run_data in para_data['runs']:
                                    run = cell_para.add_run(run_data['text'])
                                    apply_run_format(run, run_data)
                        else:
                            # 如果没有详细的段落信息，只添加文本
                            cell.text = cell_data['text']

    # 保存文档
    doc.save(output_path)
    print(f"新文档已保存到: {output_path}")


def apply_paragraph_format(paragraph, format_info):
    """应用段落格式"""
    # 使用安全的get方法访问字典，避免KeyError
    if format_info.get('alignment'):
        alignment_map = {
            'CENTER': WD_ALIGN_PARAGRAPH.CENTER,
            'CENTER (1)': WD_ALIGN_PARAGRAPH.CENTER,
            'RIGHT': WD_ALIGN_PARAGRAPH.RIGHT,
            'JUSTIFY': WD_ALIGN_PARAGRAPH.JUSTIFY,
            'LEFT': WD_ALIGN_PARAGRAPH.LEFT
        }
        alignment = format_info['alignment'].split('.')[-1] if '.' in format_info['alignment'] else format_info[
            'alignment']
        paragraph.alignment = alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)

    # 使用get方法安全访问，如果键不存在则返回None
    space_before = format_info.get('space_before')
    space_after = format_info.get('space_after')
    line_spacing = format_info.get('line_spacing')
    first_line_indent = format_info.get('first_line_indent')
    left_indent = format_info.get('left_indent')

    if space_before is not None:
        paragraph.paragraph_format.space_before = space_before

    if space_after is not None:
        paragraph.paragraph_format.space_after = space_after
    else:
        paragraph.paragraph_format.space_after = 0  # 段落后间距None会默认设置为10，需设置为0

    if line_spacing is not None:
        paragraph.paragraph_format.line_spacing = line_spacing
    if first_line_indent is not None:
        paragraph.paragraph_format.first_line_indent = first_line_indent
    if left_indent is not None:
        paragraph.paragraph_format.left_indent = left_indent


def apply_run_format(run, run_data):
    """应用运行文本格式"""
    # 使用get方法安全访问
    run.bold = run_data.get('bold', False)
    run.italic = run_data.get('italic', False)
    run.underline = run_data.get('underline', False)

    font_size = run_data.get('font_size')
    if font_size:
        run.font.size = Pt(font_size)

    font_name = run_data.get('font_name')
    if font_name:
        run.font.name = font_name
        run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    # 无font_name时采用默认字体
    if run.text.strip():
        run.font.name = 'Times New Roman'  # 设置西文字体
        run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 设置中文字体
    else:
        run.font.name = u'宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 设置中文字体

    font_color = run_data.get('font_color')
    if font_color:
        try:
            # 将十六进制颜色转换为RGB
            hex_color = font_color.lstrip('#')
            rgb = tuple(int(hex_color[i:i + 2], 16) for i in (0, 2, 4))
            run.font.color.rgb = RGBColor(*rgb)
        except:
            pass


def apply_merged_cells(table, merged_cells_info):
    """
    应用合并单元格信息
    """
    for merge_info in merged_cells_info:
        if merge_info['is_merge_start']:
            row = merge_info['row']
            col = merge_info['col']

            # 处理水平合并（gridSpan）
            if merge_info['grid_span'] > 1:
                # 合并水平单元格
                start_cell = table.rows[row].cells[col]
                end_cell = table.rows[row].cells[col + merge_info['grid_span'] - 1]
                start_cell.merge(end_cell)


def apply_table_format(table, table_format_info):
    for cell_format in table_format_info:
        row = cell_format['row']
        col = cell_format['col']
        table.cell(row, col).width = Inches(cell_format['width'])


def add_translation_paragraph(original_data, translator, language):
    """添加翻译段落"""
    new_data = []

    for item in original_data:
        new_data.append(item)
        if item['type'] == 'paragraph':
            original_text = item['text']
            if original_text.strip():
                # 在data里添加翻译内容
                # 多种语言
                for lang in language:
                    new_item = item.copy()

                    translated_text = translator.translate(original_text, language=lang)
                    new_item['text'] = translated_text
                    new_item['language'] = lang

                    # 更改run里面内容
                    new_run = item['runs'][0].copy()
                    new_run['text'] = translated_text
                    new_runs = [new_run]
                    new_item['runs'] = new_runs
                    new_data.append(new_item)

    return new_data


def add_translation_table(original_data, translator, language):
    new_data = original_data.copy()
    for item in new_data:
        if item['type'] == 'table':
            for row in item['rows']:
                for cell in row['cells']:
                    new_cell = cell.copy()
                    new_cell['paragraphs'] = []

                    for para in cell['paragraphs']:
                        original_text = para['text']
                        if original_text.strip():
                            new_cell['paragraphs'].append(para)

                            # 多种语言
                            for lang in language:
                                new_para = para.copy()

                                translated_text = translator.translate(original_text, language=lang)
                                new_para['text'] = translated_text
                                new_para['language'] = lang

                                # 更改run里面内容
                                new_run = para['runs'][0].copy()
                                new_run['text'] = translated_text
                                new_runs = [new_run]
                                new_para['runs'] = new_runs

                                new_cell['paragraphs'].append(new_para)
                    cell['paragraphs'] = new_cell['paragraphs']

    return new_data


# 预设的账号密码（在实际应用中应该使用更安全的方式存储）
VALID_ACCOUNTS = {
    "admin": "admin123",
    "user": "user123",
    "translator": "cmtech.20250924"
}


def login():
    """登录验证函数"""
    max_attempts = 3
    attempts = 0

    while attempts < max_attempts:
        print("\n" + "=" * 50)
        print("          文档翻译系统 - 登录验证")
        print("=" * 50)

        username = input("请输入用户名: ").strip()
        print("123")
        password = getpass.getpass("请输入密码: ").strip()

        # 验证账号密码
        if username in VALID_ACCOUNTS and VALID_ACCOUNTS[username] == password:
            print(f"\n✅ 登录成功！欢迎 {username}！")
            return True
        else:
            attempts += 1
            remaining_attempts = max_attempts - attempts
            print(f"\n❌ 用户名或密码错误！剩余尝试次数: {remaining_attempts}")

            if remaining_attempts > 0:
                print("请重新输入...")
            else:
                print("\n🚫 登录失败次数过多，程序退出！")
                return False

    return False


def check_license():
    """简单的许可证检查（可选功能）"""
    import datetime as dt
    expiry_date = dt.datetime(2025, 10, 31)  # 设置过期时间

    if dt.datetime.now() > expiry_date:
        print("🚫 软件许可证已过期，请联系管理员！")
        return False
    return True


if __name__ == "__main__":
    # 首先进行登录验证
    # if not login():
    #     exit()

    # 可选：进行许可证检查
    if not check_license():
        exit()

    # 获取当前时间戳
    current_time = datetime.datetime.now().strftime('%y%m%d%H%M%S')
    print(f"********************start at {current_time}********************")
    try:
        # 设置输入文件夹路径
        input_folder = input("请输入待翻译的文件目录：\n")

        # 设置输出文件夹路径（二级文件夹）
        output_folder = os.path.join(input_folder, "translate_output")

        # 创建输出文件夹（如果不存在）
        os.makedirs(output_folder, exist_ok=True)

        # 初始化翻译器
        translator = Translator()
        language = ['英文', '越南语']

        # 遍历文件夹中的所有文件
        for filename in os.listdir(input_folder):
            # 检查文件是否为Word文档（.docx格式）
            if filename.endswith('.docx') and not filename.startswith('~$'):  # 忽略临时文件
                input_file = os.path.join(input_folder, filename)

                # 生成输出文件名（原文件名+时间）
                file_base_name = os.path.splitext(filename)[0]  # 去掉扩展名
                output_filename = f"{file_base_name}_translate_{current_time}.docx"
                output_file = os.path.join(output_folder, output_filename)

                print(f"正在处理文件: {filename}")

                # 1. 读取原文档内容及格式
                content_data = get_content(input_file)
                # print(f"  总共提取了 {len(content_data)} 个内容块")

                # 2. 翻译
                after_para = add_translation_paragraph(content_data, translator, language)
                after_table = add_translation_table(after_para, translator, language)

                # 3. 创建新文档
                create_new_document(after_table, output_file)
                print(f"  已完成翻译: {output_filename}")

        print(f"所有文件处理完成！输出目录: {output_folder}")

    except Exception as e:
        print(f"处理过程中出现错误: {e}")

    print(f"********************end********************")
