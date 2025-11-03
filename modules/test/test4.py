#!/usr/bin/env python3
# -*- coding: utf-8 -*- 
"""
@Project : tools
@File    : test4.py 
@Author  : Shawn
@Date    : 2025/10/30 13:41 
@Info    : Description of this file
"""

import datetime
from docx import Document

input_file = r"D:\Code\Project\tools\data\test\test2.docx"
doc = Document(input_file)

for table in doc.tables:
    for column in table.columns:
        for cell in column.cells:
            print(cell.text)

for table in doc.tables:
    for idx, column in enumerate(table.columns):
        width = table.columns[idx].width
        print(width)


def get_cell_size(self, cell, table, row_idx, col_idx):
    """获取单元格的宽高"""
    if cell.width:
        width = cell.width.inches
    else:
        width = table.columns[col_idx].width.inches

    return width