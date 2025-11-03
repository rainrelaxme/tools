#!/usr/bin/env python3
# -*- coding:utf-8 -*-
"""
@Project: tools.py
@File   : cm_sop_translate.py
@Version:
@Author : RainRelaxMe
@Date   : 2025/9/24 02:00
@Info   : 实现word文档的翻译，包括段落、表格，并将翻译后的文本置于原文本后，且保持原文档格式。其中doc转换为doc利用了win32包，仅支持在windows系统。
"""
import logging

from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import RGBColor, Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import CT_R
from win32com import client as wc

from conf.conf import LOG_PATH
from modules.common.log import setup_logger
from modules.cm_sop_translate.template import apply_preamble_format, apply_approveTable_format

logger = setup_logger(log_dir=LOG_PATH['path'], name='logs', level=logging.INFO)


class DocumentContent:
    def __init__(self, doc, new_doc=None):
        self.source_doc = doc
        self.errors = []

    def get_content(self, doc):
        """
        读取word文件的内容，返回位置、内容、格式，记录在dataform中
        保持段落和表格在文档中的原始顺序
        """
        data = []
        # 排序获取位置
        all_elements = self.sort_element(doc)

        # 按排序后的顺序处理元素
        for element_type, position, index, text in all_elements:
            if element_type == 'paragraph':
                para = doc.paragraphs[index]
                para_data = {
                    'type': 'paragraph',
                    'index': position,
                    'element_index': index,
                    'flag': '',
                    'text': para.text,
                    'para_format': self.get_paragraph_format(para),
                    'runs': self.get_run_format(para)
                }
                data.append(para_data)

            elif element_type == 'table':
                table = doc.tables[index]
                table_data = {
                    'type': 'table',
                    'index': position,
                    'element_index': index,
                    'flag': '',
                    "rows": len(table.rows),
                    'cols': len(table.columns),
                    "table_format": {
                        "table_alignment": WD_TABLE_ALIGNMENT.CENTER,
                    },
                    "cells": self.get_table_content(table)
                }
                data.append(table_data)

        return data

    def sort_element(self, doc):
        """对元素进行排序"""
        # 先处理所有段落，记录它们在文档中的位置
        para_positions = {}
        for para_index, para in enumerate(doc.paragraphs):
            # 获取段落在XML中的位置
            para_xml = para._element
            parent = para_xml.getparent()
            if parent is not None:
                position = list(parent).index(para_xml)
                para_positions[(position, para.text)] = para_index

        # 处理所有表格，记录它们在文档中的位置
        table_positions = {}
        for table_index, table in enumerate(doc.tables):
            table_xml = table._element
            parent = table_xml.getparent()
            if parent is not None:
                position = list(parent).index(table_xml)
                table_positions[position] = table_index

        # 按照文档顺序处理所有元素
        all_elements = []

        # 收集所有段落位置
        for (pos, text), para_index in para_positions.items():
            all_elements.append(('paragraph', pos, para_index, text))

        # 收集所有表格位置
        for pos, table_index in table_positions.items():
            all_elements.append(('table', pos, table_index, None))

        # 按位置排序
        all_elements.sort(key=lambda x: x[1])

        return all_elements

    def get_table_content(self, table):
        """
        获取表格内容
        """
        # rows = []
        cells_data = []
        cells = []

        for row_index, row in enumerate(table.rows):
            # row_data = {'cells': []}
            # 由于人会拖动单元格，导致超出新产生一列。同时在上方的行内是不会显示此列的。用row.grid_cols_before记录此行前有多少列，用row.grid_cols_after记录此行后有多少列
            if row.grid_cols_before > 0 or row.grid_cols_after > 0:
                error_info = f"{table}表格前后列数不一致！！！"
                logger.error(error_info)
                self.errors.append(error_info)

            # 判断合并单元格
            for cell_index, cell in enumerate(row.cells):
                # 只能判断横向合并,还有种方法，通过合并的两个单元格相等判断
                cells.append(cell)
                grid_span = int(cell.grid_span) if cell.grid_span else 1  # 单元格的跨度，若是合并单元格则大于1
                is_merge_start = False
                if grid_span > 1:
                    is_merge_start = True
                    if len(cells) > 1 and cell == cells[-2]:  # 判断是否等于前一个单元格，即是否与其是合并单元格。
                        is_merge_start = False

                cell_content = {
                    'row': row_index,
                    'col': cell_index,
                    'width': self.get_cell_size(cell, table, row_index, cell_index),
                    # 'width': cell.width.inches,
                    'grid_span': grid_span,
                    'is_merge_start': is_merge_start,
                    "content": self.get_content(cell)
                }
                cells_data.append(cell_content)

        return cells_data

    def get_cell_size(self, cell, table, row_idx, col_idx):
        """获取单元格的宽高"""
        if cell.width:
            width = cell.width.inches
        else:
            # 如果从单元格获取不到，则获取此列的宽度
            width = table.columns[col_idx].width.inches
        return width

    def extract_pics(self, doc):
        pic_list = []
        for idx, shape in enumerate(doc.inline_shapes):
            # 获取图片的rid
            rid = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
            # 图片的名称
            name = shape._inline.docPr.name
            # 根据rid获取图片对象
            image_obj = doc.part.rels.get(rid).target_part.image
            # print("image:", image_obj)
            # print("file_name:", image_obj.filename)
            # print("扩展名:", image_obj.ext)
            # # print("图片二进制数据:", image_obj.blob)
            # print("图片二进制数据hash:", image_obj.sha1)
            # print("内容类型:", image_obj.content_type)
            #
            # print("像素宽高{}x{}:".format(image_obj.px_width, image_obj.px_height))
            # print("分辨率{} x {}".format(image_obj.horz_dpi, image_obj.vert_dpi))
            # 解析图片信息
            # width = shape.width,  # Lendth对象  可以继续调用.inches/.pt/.cm
            # height = shape.height

            # print("\n")
            pic = {
                "type": "picture",
                "index": idx,
                "rid": rid,
                "name": name
            }
            pic_list.append(pic)
        return pic_list

    def extract_shapes(self, doc):
        """ 抽取图形 """
        shape_list = []
        # 在段落中抽取图形
        for idx, para in enumerate(doc.element.body):
            for run in para:
                # 找到run
                if isinstance(run, CT_R):
                    for inner_run in run:
                        tag_name = inner_run.tag.split("}")[1]
                        if tag_name == "AlternateContent":  # lxml.etree._Element对象
                            fallback = inner_run[1]
                            pict = fallback[0]
                            shape = pict[0]  # 内部图形
                            # print("图形数据:", shape.items())
                            # textbox = list(shape) # shape内部是否有文本，需判断
                            # if textbox:
                            #     textbox = textbox[0]
                            #     # 图形内的段落文本
                            #     para_list = list(textbox[0])
                            #     print("图形内部文本:", [para.text for para in para_list])
                            spe = {
                                "type": "shape",
                                "index": idx,
                                "aid": shape.items()[0][1],
                                "name": shape.items()[1][1]
                            }
                            shape_list.append(spe)
        return shape_list

    def split_cover_body_data(self, original_data):
        """获取word中封面的内容：修订记录表格之前内容算作封面"""
        # 先判断标题是表格还是文本，如果是表格那么首页就截止到第二个表格
        cover_data = []
        body_data = []
        top_title_type = ''
        cover_table_num = 1
        for index, item in enumerate(original_data):
            if item['flag'] == 'top_title':
                top_title_type = item['type']
            if top_title_type == 'table':
                cover_table_num = 2
            if item['type'] == 'table' and item['element_index'] == cover_table_num:
                break
            # 先判断再记录
            cover_data.append(item)
        body_data = original_data[len(cover_data):]

        data = {
            "cover_data": cover_data,
            "body_data": body_data,
        }
        return data

    def get_header_content(self, doc):
        """
        获取页眉内容
        首页：first_page_header
        奇数页：header
        偶数页：even_page_header
        是否链接到前一节： is_linked_to_previous
        是否首页不同：different_first_page_header_footer
        是否奇偶页不同：doc.settings.odd_and_even_pages_header_footer
        """
        header_data = []

        for index, section in enumerate(doc.sections):
            section_data = {
                "type": 'header',
                "index": '',
                "section_index": index,
                "is_linked_to_previous": section.header.is_linked_to_previous,
                "is_different_first_page": section.different_first_page_header_footer,
                "is_different_odd_and_even": doc.settings.odd_and_even_pages_header_footer,  # 判断奇偶数页眉页脚是否相同
                "first_page_header_content": self.get_content(section.first_page_header),  # 判断首页是否相同
                "odd_page_header": self.get_content(section.header),  # 如果奇偶页相同，则全部在header里
                "even_page_header": self.get_content(section.even_page_header),
            }
            header_data.append(section_data)
        return header_data

    def get_footer_content(self, doc):
        """
        获取页脚内容
        首页：first_page_footer
        奇数页：footer
        偶数页：even_page_footer
        是否链接到前一节： is_linked_to_previous
        是否首页不同：different_first_page_header_footer
        是否奇偶页不同：doc.settings.odd_and_even_pages_header_footer
        """
        footer_data = []

        for index, section in enumerate(doc.sections):
            section_data = {
                "type": 'footer',
                "index": '',
                "section_index": index,
                "is_linked_to_previous": section.footer.is_linked_to_previous,
                "is_different_first_page": section.different_first_page_header_footer,
                "is_different_odd_and_even": doc.settings.odd_and_even_pages_header_footer,  # 判断奇偶数页眉页脚是否相同
                "first_page_footer_content": self.get_content(section.first_page_footer),  # 判断首页是否相同
                "odd_page_footer": self.get_content(section.footer),  # 如果奇偶页相同，则全部在header里
                "even_page_footer": self.get_content(section.even_page_footer),
            }
            footer_data.append(section_data)
        return footer_data

    def get_paragraph_format(self, paragraph):
        """提取段落格式信息"""
        return {
            'style': paragraph.style.name if paragraph.style else 'Normal',
            'alignment': str(paragraph.alignment) if paragraph.alignment else None,
            'space_before': paragraph.paragraph_format.space_before,
            'space_after': paragraph.paragraph_format.space_after,
            'line_spacing': paragraph.paragraph_format.line_spacing,
            'first_line_indent': paragraph.paragraph_format.first_line_indent,
            'left_indent': paragraph.paragraph_format.left_indent
        }

    def get_run_format(self, paragraph):
        # 获取段落格式
        runs_data = []
        for run in paragraph.runs:
            run_format = {
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_size': run.font.size.pt if run.font.size else None,
                'font_name': run.font.name,
                'font_color': run.font.color.rgb if run.font.color and run.font.color.rgb else None,
                'font_color_theme': run.font.color.theme_color if run.font.color else None
            }

            runs_data.append(run_format)

        return runs_data

    def flag_title(self, original_data):
        """
        标记文件大标题：第一个非空内容，表格、段落都可以。
        """
        # 若是标题是用表格做的，还要处理
        title_text = ''
        for index, item in enumerate(original_data):
            if item['type'] == 'paragraph' and item['text'].strip():
                original_data[index]['flag'] = 'top_title'
                # title_text = item['text'].strip()
                break
            if item['type'] == 'table' and (len(item['rows']) == 1 and item['cols'] == 1):
                # item['rows'][0]['cells'][0]['text']
                original_data[index]['flag'] = 'top_title'
                title_text = item['rows'][0]['cells'][0]['text'].strip()
                break
        return original_data

    def flag_preamble(self, original_data):
        """
        标记文档的头信息，如文件编号、版本、制定部门、制定日期、总页次。
        """
        # 找到文件头信息的位置，从标题后第一个有内容的段落
        # 先只考虑大标题是段落的形式
        for index, item in enumerate(original_data):
            if item['flag'] == 'top_title':
                title_end_index = index
            if item['type'] == 'table':
                preamble_end_index = index
        cut_title_data = original_data[title_end_index + 1:preamble_end_index]

        for index, item in enumerate(cut_title_data):
            if item['type'] == 'paragraph' and item['text'].strip():
                cut_title_data[index]['flag'] = 'preamble'
        return original_data

    def flag_approveTable(self, original_data):
        """
        标记文档封面的审批表格
        """
        for item in original_data:
            if item['type'] == 'table' and item['element_index'] == 0:
                item['flag'] = 'approve'
        return original_data

    def flag_main_text(self, original_data):
        """
        标记出来表格的主文本所在的单元格，因为修订记录和主文本原文件中是同一个表格，要切分开。
        采用的方法是第二个表格中的第一次合并单元格，此单元格前应空行很多行，而且正文开头是“目的”？
        """
        for item in original_data:
            # 第二个表格，而且表格第一格中文字内容是“版本”，就认为是修订记录表格

            if item['type'] == 'table' and item['element_index'] == 1 and item["cells"][0]["content"][0][
                'text'].strip() == "版本":
                cols = item['cols']
                item['flag'] = 'revision_record'
                for cell in item['cells']:
                    # 合并的单元格，且合并了所有列，且第一行的正文包含“目的”，就认为是正文主体
                    if cell["grid_span"] > 1 and cell["grid_span"] == cols:  # 全部行合并
                        cell["flag"] = "main_text"
        return original_data


def set_paper_size_format(doc):
    """
    设置纸张大小和页边距:A4,页边距2/1.1/2/2
    """
    # 获取第一个节（默认存在）
    section = doc.sections[0]

    # 设置纸张方向（纵向或横向）
    section.orientation = WD_ORIENTATION.PORTRAIT  # PORTRAIT: 纵向, LANDSCAPE: 横向

    # 设置纸张大小（A4）
    section.page_width = Cm(21)  # A4宽度 21cm
    section.page_height = Cm(29.7)  # A4高度 29.7cm

    # 设置页边距（上下左右）
    section.top_margin = Cm(2)  # 上边距
    section.bottom_margin = Cm(1.1)  # 下边距
    section.left_margin = Cm(2)  # 左边距
    section.right_margin = Cm(2)  # 右边距


def add_cover(doc, data):
    """添加封面"""
    for item in data:
        if item['type'] == 'paragraph':
            # 创建新段落
            paragraph = doc.add_paragraph()

            # 应用文件头信息格式
            if item['flag'] == 'preamble':
                apply_preamble_format(paragraph, item)
            else:
                # 应用段落格式
                apply_paragraph_format(paragraph, item['para_format'])
                # 应用运行文本、格式
                for run_data in item['runs']:
                    run = paragraph.add_run(run_data['text'])
                    apply_run_format(run, run_data)
        if item['type'] == 'table':
            table = doc.add_table(rows=item['rows'], cols=item['cols'])
            table.style = 'Table Grid'

            # 设置审批表格样式
            # if item['flag'] == 'approve':
            #     apply_approveTable_format(table)

            # 设置表格样式
            apply_table_format(table, item)
            # 应用表格内容
            for cell_idx, cell_data in enumerate(item['cells']):
                if (cell_data['grid_span'] > 1 and cell_data['is_merge_start']) or cell_data['grid_span'] == 1:
                    cell = table.rows[cell_data["row"]].cells[cell_data["col"]]
                    # 清空默认段落
                    for paragraph in cell.paragraphs:
                        p = paragraph._element
                        p.getparent().remove(p)
                    # 添加内容到单元格
                    if cell_data['content']:
                        for content in cell_data['content']:
                            if content['type'] == 'paragraph':
                                cell_para = cell.add_paragraph()
                                apply_paragraph_format(cell_para, content['para_format'])

                                for run_data in content['runs']:
                                    run = cell_para.add_run(run_data['text'])
                                    apply_run_format(run, run_data)

            # 应用表格内容
            # for row_idx, row_data in enumerate(item['rows']):
            #     for cell_idx, cell in enumerate(row_data['cells']):
            #         if (cell['grid_span'] > 1 and cell['is_merge_start']) or cell['grid_span'] == 1:
            #             cell = table.rows[row_idx].cells[cell_idx]
            #
            #             # 清空默认段落
            #             for paragraph in cell.paragraphs:
            #                 p = paragraph._element
            #                 p.getparent().remove(p)
            #
            #             # 添加内容到单元格
            #             if cell['content']:
            #                 for para_data in cell['content']:
            #                     cell_para = cell.add_paragraph()
            #                     apply_paragraph_format(cell_para, para_data['para_format'])
            #
            #                     for run_data in para_data['runs']:
            #                         run = cell_para.add_run(run_data['text'])
            #                         apply_run_format(run, run_data)
            #             else:
            #                 # 如果没有详细的段落信息，只添加文本
            #                 cell.text = cell['text']


def add_content(block, data):
    """对块（doc、section、header、footer）中添加内容"""
    for index, item in enumerate(data):
        if item['type'] == 'paragraph':
            # 创建新段落
            paragraph = block.add_paragraph()

            # 应用文件头信息格式
            if item['flag'] == 'preamble':
                apply_preamble_format(paragraph, item)
            else:
                # 应用段落格式
                apply_paragraph_format(paragraph, item['para_format'])
                # 应用运行文本、格式
                for run_data in item['runs']:
                    run = paragraph.add_run(run_data['text'])
                    apply_run_format(run, run_data)
        if item['type'] == 'table':
            # 创建表格
            # # 在表格前添加分页符
            if item['element_index'] >= 0:
                page_break_para = block.add_paragraph()
                run = page_break_para.add_run()
                run.add_break(WD_BREAK.PAGE)

            table = block.add_table(rows=item['rows'], cols=item['cols'])
            table.style = 'Table Grid'

            # 设置审批表格样式
            if item['flag'] == 'approve':
                apply_approveTable_format(table)

            # 设置表格样式
            apply_table_format(table, item)

            # 应用表格内容
            for cell_idx, cell_data in enumerate(item['cells']):
                if (cell_data['grid_span'] > 1 and cell_data['is_merge_start']) or cell_data['grid_span'] == 1:
                    cell = table.rows[cell_data["row"]].cells[cell_data["col"]]
                    # 清空默认段落
                    for paragraph in cell.paragraphs:
                        p = paragraph._element
                        p.getparent().remove(p)
                    # 添加内容到单元格
                    if cell_data['content']:
                        for content in cell_data['content']:
                            if content['type'] == 'paragraph':
                                cell_para = cell.add_paragraph()
                                apply_paragraph_format(cell_para, content['para_format'])

                                for run_data in content['runs']:
                                    run = cell_para.add_run(run_data['text'])
                                    apply_run_format(run, run_data)

                            # 表格内嵌表格
                            if content['type'] == 'table':
                                sub_table = cell.add_table(rows=content['rows'], cols=content['cols'])
                                sub_table.style = 'Table Grid'

                                # 设置表格样式
                                apply_table_format(sub_table, content)

                                # 应用表格内容
                                for sub_cell_idx, sub_cell_data in enumerate(content['cells']):
                                    if (sub_cell_data['grid_span'] > 1 and sub_cell_data['is_merge_start']) or \
                                            sub_cell_data['grid_span'] == 1:
                                        sub_cell = sub_table.rows[sub_cell_data["row"]].cells[sub_cell_data["col"]]
                                        # 清空默认段落
                                        for sub_paragraph in sub_cell.paragraphs:
                                            sub_p = sub_paragraph._element
                                            sub_p.getparent().remove(sub_p)
                                        # 添加内容到单元格
                                        if sub_cell_data['content']:
                                            for sub_content in sub_cell_data['content']:
                                                if sub_content['type'] == 'paragraph':
                                                    sub_cell_para = sub_cell.add_paragraph()
                                                    apply_paragraph_format(sub_cell_para, sub_content['para_format'])

                                                    for run_data in sub_content['runs']:
                                                        run = sub_cell_para.add_run(run_data['text'])
                                                        apply_run_format(run, run_data)


def add_header(doc, header_data):
    """
    添加页眉
    首页：first_page_header
    奇数页：header
    偶数页：even_page_header
    是否链接到前一节： is_linked_to_previous
    是否首页不同：different_first_page_header_footer
    是否奇偶页不同：doc.settings.odd_and_even_pages_header_footer
    """
    for index, item in enumerate(header_data):
        # 创建节
        section = doc.sections[item['section_index']]
        # 是否链接到前一节
        section.header.is_linked_to_previous = item['is_linked_to_previous']
        # 设置奇偶页是否相同
        doc.settings.odd_and_even_pages_header_footer = item['is_different_odd_and_even']
        # 设置首页是否相同
        section.different_first_page_header_footer = item['is_different_first_page']

        # 填充内容
        if item['first_page_header_content']:
            add_content(section.first_page_header, item['first_page_header_content'])
        if item['odd_page_header']:
            add_content(section.header, item['odd_page_header'])
        if item['even_page_header']:
            add_content(section.even_page_header, item['even_page_header'])


def add_footer(doc, footer_data):
    """
    添加页脚
    首页：first_page_footer
    奇数页：footer
    偶数页：even_page_footer
    是否链接到前一节： is_linked_to_previous
    是否首页不同：different_first_page_header_footer
    是否奇偶页不同：doc.settings.odd_and_even_pages_header_footer
    """
    for index, item in enumerate(footer_data):
        # 创建节
        section = doc.sections[item['section_index']]
        # 是否链接到前一节
        section.header.is_linked_to_previous = item['is_linked_to_previous']
        # 设置奇偶页是否相同
        doc.settings.odd_and_even_pages_header_footer = item['is_different_odd_and_even']
        # 设置首页页脚不同
        section.different_first_page_header_footer = item['is_different_first_page']

        # 填充内容
        if item['first_page_footer_content']:
            add_content(section.first_page_footer, item['first_page_footer_content'])
        if item['odd_page_footer']:
            add_content(section.footer, item['odd_page_footer'])
        if item['even_page_footer']:
            add_content(section.even_page_header, item['even_page_footer'])


def apply_paragraph_format(paragraph, format_info):
    """应用段落格式"""
    # 使用安全的get方法访问字典，避免KeyError
    if format_info.get('alignment'):
        alignment_map = {
            'CENTER': WD_ALIGN_PARAGRAPH.CENTER,
            'CENTER (1)': WD_ALIGN_PARAGRAPH.CENTER,
            'RIGHT': WD_ALIGN_PARAGRAPH.RIGHT,
            'JUSTIFY': WD_ALIGN_PARAGRAPH.JUSTIFY,
            'LEFT': WD_ALIGN_PARAGRAPH.LEFT
        }
        alignment = format_info['alignment'].split('.')[-1] if '.' in format_info['alignment'] else format_info[
            'alignment']
        paragraph.alignment = alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)

    # 使用get方法安全访问，如果键不存在则返回None
    space_before = format_info.get('space_before')
    space_after = format_info.get('space_after')
    line_spacing = format_info.get('line_spacing')
    first_line_indent = format_info.get('first_line_indent')
    left_indent = format_info.get('left_indent')

    if space_before is not None:
        paragraph.paragraph_format.space_before = space_before

    if space_after is not None:
        paragraph.paragraph_format.space_after = space_after
    else:
        paragraph.paragraph_format.space_after = 0  # 段落后间距None会默认设置为10，需设置为0

    if line_spacing is not None:
        paragraph.paragraph_format.line_spacing = line_spacing
    if first_line_indent is not None:
        paragraph.paragraph_format.first_line_indent = first_line_indent
    if left_indent is not None:
        paragraph.paragraph_format.left_indent = left_indent


def apply_run_format(run, run_data):
    """应用运行文本格式"""
    # 使用get方法安全访问
    run.bold = run_data.get('bold', False)
    run.italic = run_data.get('italic', False)
    run.underline = run_data.get('underline', False)

    font_size = run_data.get('font_size')
    if font_size:
        run.font.size = Pt(font_size)

    font_name = run_data.get('font_name')
    if font_name:
        run.font.name = font_name
        run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    # 无font_name时采用默认字体
    if run.text.strip():
        run.font.name = 'Times New Roman'  # 设置西文字体
        run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 设置中文字体
    else:
        run.font.name = u'宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 设置中文字体

    font_color = run_data.get('font_color')
    if font_color:
        try:
            # 将十六进制颜色转换为RGB
            # hex_color = font_color.lstrip('#')
            # rgb = tuple(int(hex_color[i:i + 2], 16) for i in (0, 2, 4))
            # run.font.color.rgb = RGBColor(*rgb)
            run.font.color.rgb = RGBColor(*font_color)
        except:
            pass


def apply_table_format(table, table_data):
    """
    处理表格样式
    """
    # 自动调整列宽：禁用
    table.autofit = False
    # 表格对齐方式，非内容对齐方式
    table.alignment = table_data["table_format"]['table_alignment']

    # 单元格样式
    for cell in table_data["cells"]:
        row = cell['row']
        col = cell['col']
        # 设置表格宽度
        if cell.get("width"):
            table.cell(row, col).width = Inches(cell['width'])
        table.rows[row].height = Cm(1.3)

        # 合并单元格
        if cell["is_merge_start"]:
            # 处理水平合并（gridSpan）
            if cell['grid_span'] > 1:
                start_cell = table.rows[row].cells[col]
                end_cell = table.rows[row].cells[col + cell['grid_span'] - 1]
                start_cell.merge(end_cell)


def add_paragraph_translation(original_data, translator, language):
    """添加翻译段落"""
    new_data = []

    for item in original_data:
        new_data.append(item)
        if item['type'] == 'paragraph':
            original_text = item['text']
            if original_text.strip():
                # 在data里添加翻译内容
                # 多种语言
                for lang in language:
                    new_item = item.copy()

                    translated_text = translator.translate(original_text, language=lang, display=True)
                    # translated_text = original_text
                    new_item['text'] = translated_text
                    new_item['language'] = lang

                    # 更改run里面内容
                    new_run = item['runs'][0].copy()
                    new_run['text'] = translated_text
                    new_runs = [new_run]
                    new_item['runs'] = new_runs
                    new_data.append(new_item)

    return new_data


def add_table_translation(original_data, translator, language):
    new_data = original_data.copy()
    new_data2 = []
    for index, item in enumerate(original_data):
        if item['type'] == 'table':
            for idx, cell in enumerate(item['cells']):
                if (cell['grid_span'] > 1 and cell['is_merge_start']) or cell['grid_span'] == 1:
                    after_para = add_paragraph_translation(cell["content"], translator, language)
                    after_table = add_table_translation(after_para, translator, language)

                    new_data[index]['cells'][idx]["content"] = after_table

            #     new_item = item.copy()
            #
            # for row in item['rows']:
            #     for cell in row['cells']:
            #         new_cell = cell.copy()
            #         new_cell['content'] = []
            #
            #         for para in cell['content']:
            #             if para['type'] == 'paragraph':
            #                 original_text = para['text']
            #                 if original_text.strip():
            #                     if (cell['grid_span'] > 1 and cell['is_merge_start']) or cell['grid_span'] == 1:
            #                         new_cell['content'].append(para)
            #
            #                         # 多种语言
            #                         for lang in language:
            #                             new_para = para.copy()
            #
            #                             translated_text = translator.translate(original_text, language=lang,
            #                                                                    display=True)
            #                             # translated_text = original_text
            #                             new_para['text'] = translated_text
            #                             new_para['language'] = lang
            #
            #                             # 更改run里面内容
            #                             new_run = para['runs'][0].copy()
            #                             new_run['text'] = translated_text
            #                             new_runs = [new_run]
            #                             new_para['runs'] = new_runs
            #
            #                             new_cell['content'].append(new_para)
            #                 if para['type'] == 'table':
            #                     for row in item['rows']:
            #                         for cell in row['cells']:
            #                             new_cell = cell.copy()
            #                             new_cell['content'] = []
            #
            #                             for para in cell['content']:
            #                                 if para['type'] == 'paragraph':
            #                                     original_text = para['text']
            #                                     if original_text.strip():
            #                                         if (cell['grid_span'] > 1 and cell['is_merge_start']) or cell['grid_span'] == 1:
            #                                             new_cell['content'].append(para)
            #
            #                                             # 多种语言
            #                                             for lang in language:
            #                                                 new_para = para.copy()
            #
            #                                                 translated_text = translator.translate(original_text,
            #                                                                                        language=lang,
            #                                                                                        display=True)
            #                                                 # translated_text = original_text
            #                                                 new_para['text'] = translated_text
            #                                                 new_para['language'] = lang
            #
            #                                                 # 更改run里面内容
            #                                                 new_run = para['runs'][0].copy()
            #                                                 new_run['text'] = translated_text
            #                                                 new_runs = [new_run]
            #                                                 new_para['runs'] = new_runs
            #
            #                                                 new_cell['content'].append(new_para)
            #                         cell['content'] = new_cell['content']
            #
            #         cell['content'] = new_cell['content']
    return new_data


def add_cover_translation(original_data, translator, language):
    """封面的'文件编号：C2GM-Z13-000'这种需要通过切分分号来分开翻译"""
    # 获取封面内容
    for index, item in enumerate(original_data):
        if item['flag'] == 'top_title':
            title_end_index = index + 1
        if item['flag'] == 'preamble':
            preamble_end_index = index + 1

    # 1. '文件编号'之前的正常翻译
    cover_data_1 = original_data[:title_end_index]
    after_para_1 = add_paragraph_translation(cover_data_1, translator, language)
    after_table_1 = add_table_translation(after_para_1, translator, language)

    # 2. 文件头信息内容：文件编号之后到审批表格前
    cover_data_2 = original_data[title_end_index:preamble_end_index]
    translated_cover_data_2 = []
    for item in cover_data_2:
        translated_cover_data_2.append(item)
        if item['type'] == 'paragraph' and item['flag'] == 'preamble':
            original_text = item['text']

            # 根据冒号拆分段落
            split_text = []
            if ":" in original_text:
                split_text = original_text.split(":", 1)  # 1代表只分割第一个冒号
            elif "：" in original_text:
                split_text = original_text.split("：", 1)  # 1代表只分割第一个冒号

            if split_text:
                for lang in language:
                    new_item = item.copy()

                    translated_text = ''
                    for part in split_text:
                        part_translated = translator.translate(part, language=lang, display=True)
                        translated_text = translated_text + part_translated + '：'
                    translated_text = translated_text[:-1]
                    # print(translated_text)
                    new_item['text'] = translated_text
                    new_item['language'] = lang

                    # 更改run里面内容
                    new_run = item['runs'][0].copy()
                    new_run['text'] = translated_text
                    new_runs = [new_run]
                    new_item['runs'] = new_runs
                    translated_cover_data_2.append(new_item)

    # 3. 审批表格翻译
    cover_data_3 = original_data[preamble_end_index:]
    after_para_3 = add_paragraph_translation(cover_data_3, translator, language)
    after_table_3 = add_table_translation(after_para_3, translator, language)

    # 4. 汇总
    new_cover_data = []
    for item in after_table_1:
        new_cover_data.append(item)
    for item in translated_cover_data_2:
        new_cover_data.append(item)
    for item in after_table_3:
        new_cover_data.append(item)

    return new_cover_data


def add_header_translation(original_data, translator, language):
    """页眉翻译，只有最顶部的“秘密”是不产生新的段落"""
    new_data = []
    for item in original_data:
        item_translation = {
            'type': item['type'],
            'index': item['index'],
            'section_index': item['section_index'],
            "is_linked_to_previous": item['is_linked_to_previous'],
            "is_different_first_page": item['is_different_first_page'],
            "is_different_odd_and_even": item['is_different_odd_and_even'],
            "first_page_header_content": add_table_translation(
                add_paragraph_translation(item['first_page_header_content'], translator, language), translator,
                language),
            'odd_page_header': add_table_translation(
                add_paragraph_translation(item['odd_page_header'], translator, language), translator, language),
            'even_page_header': add_table_translation(
                add_paragraph_translation(item['even_page_header'], translator, language), translator, language)
        }
        # after_para = add_paragraph_translation(item['content'], translator, language)
        # after_table = add_table_translation(after_para, translator, language)
        # for i in after_table:
        #     item_translation['content'].append(i)
        new_data.append(item_translation)
    return new_data


def add_footer_translation(original_data, translator, language):
    """页脚翻译"""
    new_data = []
    for index, item in enumerate(original_data):
        item_translation = {
            'type': item['type'],
            'index': item['index'],
            'section_index': item['section_index'],
            "is_linked_to_previous": item['is_linked_to_previous'],
            "is_different_first_page": item['is_different_first_page'],
            "is_different_odd_and_even": item['is_different_odd_and_even'],
            "first_page_footer_content": add_table_translation(
                add_paragraph_translation(item['first_page_footer_content'], translator, language), translator,
                language),
            'odd_page_footer': add_table_translation(
                add_paragraph_translation(item['odd_page_footer'], translator, language), translator, language),
            'even_page_footer': add_table_translation(
                add_paragraph_translation(item['even_page_footer'], translator, language), translator, language)
        }
        new_data.append(item_translation)
    return new_data


def doc_to_docx(doc_path, docx_path=None):
    """
    将doc文件转换为docx格式
    """
    # 如果未指定输出路径，自动生成
    if docx_path is None:
        docx_path = doc_path.replace('.doc', '.docx')

    # 启动Word应用程序
    word = wc.Dispatch('Word.Application')
    word.Visible = False  # 不显示Word界面

    try:
        # 打开doc文档
        doc = word.Documents.Open(doc_path)
        # 另存为docx格式
        doc.SaveAs(docx_path, FileFormat=16)  # 16表示docx格式
        doc.Close()
        logger.info(f"转换成功: {doc_path} -> {docx_path}")
        return docx_path
    except Exception as e:
        logger.error(f"转换失败: {e}")
        return False
    finally:
        word.Quit()
