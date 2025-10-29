# -*- coding:utf-8 -*-
# @Time   : 2023-11-08
# @Author : Carl_DJ
import io

from docx import Document
from PIL import Image
import os

from docx.shared import Cm

input_file = r"D:\Code\Project\tools\data\test\test.docx"
doc = Document(input_file)
def extract_pics(block):

    for idx, shape in enumerate(block.paragraphs):
        rid = shape._inline





# 创建一个新的Word文档
new_doc = Document()

# 获取图像目录中的所有文件名
image_dir = r"D:\Code\Project\tools\data\images"
images = [f for f in os.listdir(image_dir) if os.path.isfile(os.path.join(image_dir, f))]

# 将每个图像插入到文档中
row_num = 0
col_num = 0
table = doc.add_table(rows=2, cols=2)
cell = table.rows[0].cells[0]
for i, image_name in enumerate(images):
    # 打开图像
    img = Image.open(os.path.join(image_dir, image_name))

    # 将图像转换为BytesIO对象
    byte_arr = io.BytesIO()
    img.save(byte_arr, format='PNG')
    byte_arr.seek(0)

    # 插入图像
    run = cell.add_paragraph().add_run()
    inline_shape = run.add_picture(byte_arr, width=Cm(2), height=Cm(2))

    col_num += 1

    # 如果当前行已经满了（即插入了4张图片），则开始新的一行
    if col_num == 4:
        row_num += 1
        col_num = 0

# 保存文档
new_doc.save('output.docx')


