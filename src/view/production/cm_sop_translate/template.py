#!/usr/bin/env python3
# -*- coding: utf-8 -*- 
"""
@Project : tools
@File    : template.py 
@Author  : Shawn
@Date    : 2025/9/30 14:15 
@Info    : Description of this file
"""

import datetime
import json
import os
import win32com.client as wc
from win32com.client import constants
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_TAB_ALIGNMENT, WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Cm, RGBColor

from config.template_content import FOOTER_FORMAT, HEADER_FORMAT


def apply_template(body_data, header_data=None, footer_data=None, cover_data=None):
    """应用模板内容，修改data"""
    # after_header = apply_header_template(header_data)
    # after_footer = apply_footer_template(footer_data)
    after_cover = apply_cover_template(cover_data)
    after_body = apply_main_text_template(body_data)
    new_data = {
        "header": header_data,
        "footer": footer_data,
        "cover": after_cover,
        "body": after_body,
    }
    return new_data


def apply_header_template(doc, header_data):
    """应用页眉模板，修改data"""
    # 如果没有此内容，则不应用模板。
    if header_data is None:
        return None

    # 设置奇偶页是否相同
    doc.settings.odd_and_even_pages_header_footer = False

    for index, item in enumerate(header_data):
        # 创建节
        section = doc.sections[item['section_index']]

        # 设置首页是否相同
        section.different_first_page_header_footer = True
        # 填充内容
        if item['first_page_header_content']:
            add_content(section.first_page_header, item['first_page_header_content'])
        if item['odd_page_header']:
            add_content(section.header, item['odd_page_header'])
    return True


def apply_footer_template(doc, footer_data):
    """应用页眉模板，修改data"""
    # 如果没有此内容，则不应用模板。
    if footer_data is None:
        return None

    # 设置奇偶页是否相同
    doc.settings.odd_and_even_pages_header_footer = False

    for index, item in enumerate(footer_data):
        # 创建节
        section = doc.sections[item['section_index']]

        # 设置首页页脚不同
        section.different_first_page_header_footer = True
        # 填充内容
        if item['first_page_footer_content']:
            add_content(section.first_page_footer, item['first_page_footer_content'])
        if item['odd_page_footer']:
            add_content(section.footer, item['odd_page_footer'])
    return True


def set_page_no():
    """设置页面的页码"""
    doc_app = wc.gencache.EnsureDispatch('Word.Application')
    doc_app.Visible = False


def apply_cover_template(cover_data=None):
    """应用封面模板，修改data"""
    # 如果没有此内容，则不应用模板。
    if cover_data is None:
        return None

    new_cover_data = []
    para_data = {
        'type': 'paragraph',
        'index': None,
        'element_index': None,
        'flag': '',
        'text': '',
        'para_format': {
            'style': 'Normal',
            'alignment': None,
            'space_before': None,
            'space_after': None,
            'line_spacing': 2.0,
            'first_line_indent': None,
            'left_indent': None
        },
        'runs': []
    }
    table_data = {
        'type': 'table',
        'index': None,
        'element_index': None,
        'flag': None,
        'rows': None,
        'cols': None,
        "table_format": {
            "table_alignment": WD_TABLE_ALIGNMENT.CENTER,
        },
        "cells": []
    }

    # 1. 第一行空格
    para_1 = para_data.copy()
    para_1['index'] = 0
    para_1['element_index'] = 0
    new_cover_data.append(para_1)

    # 2. 大标题
    table_1 = table_data.copy()
    table_1['index'] = 1
    table_1['element_index'] = 0
    table_1['flag'] = 'top_title'
    table_1['rows'] = 1
    table_1['cols'] = 1
    table_1["cells"] = []
    cell = {
        "row": 0,
        "col": 0,
        "width": "",
        "grid_span": 1,
        "is_merge_start": False,
        "content": []
    }
    para_format = {
        'style': 'Normal',
        'alignment': 'CENTER (1)',
        'space_before': None,
        'space_after': None,
        'line_spacing': None,
        'first_line_indent': None,
        'left_indent': None
    }
    for item in cover_data:
        if item["flag"] == 'top_title':
            content = {
                "type": "paragraph",
                "index": 0,
                "element_index": 0,
                "flag": "",
                "text": item["text"],
                'para_format': para_format,
                'runs': item['runs'] if 'runs' in item else []
            }
            cell['content'].append(content)
    table_1["cells"].append(cell)
    new_cover_data.append(table_1)

    # 3. 接空行
    para_2 = para_data.copy()
    para_2['index'] = 2
    para_2['element_index'] = 1
    new_cover_data.append(para_2)

    # 4. 文件头信息：文档编号等
    index = len(new_cover_data)
    para_index = 2
    i = 0
    for item in cover_data:
        if item['flag'] == 'preamble':
            item['index'] = index
            item['element_index'] = para_index
            new_cover_data.append(item)
            index += 1
            para_index += 1
            i += 1
            if i < 3:  # 翻译语言的种类
                continue

            # 加一个空行
            new_blank = para_data.copy()
            new_blank['index'] = index + 1
            new_blank['element_index'] = para_index + 1
            new_cover_data.append(new_blank)
            index += 1
            para_index += 1
            i = 0

    # 5. 审批表格
    for item in cover_data:
        if item['type'] == 'table' and item['flag'] == 'approve':
            item['index'] = len(new_cover_data) + 1
            item['element_index'] = 1
            for cell in item["cells"]:
                cell["width"] = 1.8
                for para in cell["content"]:
                    para["para_format"]["line_spacing"] = Pt(20)
                    para["para_format"]["first_line_indent"] = None
                    para["para_format"]["alignment"] = "CENTER (1)"

            new_cover_data.append(item)

    return new_cover_data


def apply_main_text_template(body_data=None):
    """先拆分修订记录和正文表格，然后修改数据，使其正文单起一个表格"""
    # 如果没有此内容，则不应用模板。
    if body_data is None:
        return None

    for item in body_data:
        if item['type'] == 'table' and item["flag"] == "revision_record":
            # 正文主体的表格（list）
            main_text_cells = []
            for cell in item["cells"]:
                if cell.get("flag") == "main_text":
                    main_text_cells.append(cell)

            # 修订记录的表格（list）
            revision_cells = [i for i in item["cells"] if i not in main_text_cells]
            # 保留12行
            is_allow_cut = True
            line_num = 10
            if len(revision_cells) >= line_num:
                for cell in revision_cells[line_num*item["cols"]:]:
                    for content in cell["content"]:
                        if content["text"].strip() != "":
                            is_allow_cut = False
                            break
            if is_allow_cut:
                revision_cells = revision_cells[:line_num*item["cols"]]

            # 生成新的body_data
            item["cells"] = revision_cells
            item["rows"] = int(len(revision_cells)/item["cols"])

            if len(main_text_cells) > 0:
                main_text_cell_data = main_text_cells[0]
                main_text_cell_data["grid_span"] = 1
                main_text_cell_data["is_merge_start"] = False
                main_text_cell_data["row"] = 0
                main_text_cell_data["col"] = 0
                # width = 0
                # for cell1 in main_text_cells:
                #     width += cell1["width"]
                # main_text_cell_data["width"] = width
                main_text_data = {
                    "type": "table",
                    "index": int(item["index"]) + 1,
                    "element_index": int(item["element_index"]) + 1,
                    "flag": "main_text",
                    "rows": 1,
                    "cols": 1,
                    "table_format": item["table_format"],
                    "cells": [main_text_cell_data]
                }
                # 插入到当前表格后面，即修订记录之后
                body_data.insert(body_data.index(item)+1, main_text_data)

    return body_data


def apply_preamble_format(paragraph, preamble_data):
    """
    应用文件头信息格式，通过制表符对齐，生成的格式
    """
    # 清除默认制表位
    paragraph.paragraph_format.tab_stops.clear_all()

    # 固定行距20磅
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY  # 固定行距
    paragraph_format.line_spacing = Pt(20)
    paragraph_format.space_after = 0

    # 添加自定义制表位
    tab_stop1 = paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(1.7),
        WD_TAB_ALIGNMENT.LEFT
    )
    tab_stop2 = paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(3.5),
        WD_TAB_ALIGNMENT.LEFT
    )

    # 使用制表位
    split_text = preamble_data['text'].split('：')
    paragraph.add_run("\t")  # 制表符1
    run1 = paragraph.add_run(split_text[0] + '：')
    run1.font.bold = True
    run1.font.size = Pt(16.0)
    run1.font.name = u'宋体'
    run1.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 设置中文字体

    paragraph.add_run("\t")  # 制表符2
    run2 = paragraph.add_run(split_text[1].strip())
    run2.font.size = Pt(16.0)
    run2.font.name = 'Times New Roman'  # 设置西文字体
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 设置中文字体


def apply_approveTable_format(table):
    """
    应用审批表格格式，生成的格式
    """
    for i in range(len(table.rows)):
        table.rows[i].height = Cm(2.5)

    for i in range(len(table.rows)):
        for j in range(len(table.columns)):
            table.cell(i, j).width = Cm(5)


def apply_header_format(doc, header_data):
    """应用页眉，生成的格式"""
    # 设置奇偶页是否相同
    doc.settings.odd_and_even_pages_header_footer = False

    # 创建节
    section = doc.sections[0]
    # 设置首页不同
    section.different_first_page_header_footer = True

    # 首页页眉
    # 先清除内容
    for para in section.first_page_header.paragraphs:
        p = para._element
        p.getparent().remove(p)
    # 内容
    first_para1 = section.first_page_header.add_paragraph()
    first_para1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    first_para1.paragraph_format.space_after = 0
    first_run1 = first_para1.add_run(HEADER_FORMAT["tag"])
    first_run1.font.size = Pt(22.0)
    first_run1.font.name = 'Times New Roman'  # 设置西文字体
    first_run1.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 设置中文字体
    # 横线
    first_table1 = section.first_page_header.add_table(rows=1, cols=1, width=Cm(18))
    first_table1.style = "Table Grid"
    first_table1.autofit = False
    first_table1.alignment = WD_TAB_ALIGNMENT.CENTER

    first_cell = first_table1.cell(0, 0)
    set_cell_border(
        first_cell,
        start={"sz": 12, "val": "none"},
        end={"sz": 12, "val": "none"},
        top={"sz": 12, "val": "none"},
        bottom={"sz": 5, "val": "single", "color": "#000000"},
    )

    # 非首页页眉
    # 先清除内容
    for para in section.header.paragraphs:
        p = para._element
        p.getparent().remove(p)
    # 内容
    para1 = section.header.add_paragraph()
    para1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    para1.paragraph_format.space_after = 0
    run1 = para1.add_run(HEADER_FORMAT["tag"])
    run1.font.size = Pt(18.0)
    run1.font.name = 'Times New Roman'  # 设置西文字体
    run1.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 设置中文字体

    # 表格
    table = section.header.add_table(rows=2, cols=5, width=section.page_width)
    table.style = "Table Grid"
    table.autofit = False
    table.alignment = WD_TAB_ALIGNMENT.CENTER
    for item in header_data[0]["odd_page_header"]:
        # 页眉的表格
        if item["type"] == "table":
            for cell in item["cells"]:
                new_cell = table.cell(cell["row"], cell["col"])
                # 先清除段落内容
                for para in new_cell.paragraphs:
                    p = para._element
                    p.getparent().remove(p)
                new_cell.width = Inches(cell["width"])
                for para in cell["content"]:
                    new_para = new_cell.add_paragraph()
                    run = new_para.add_run(para["text"])
                    run.font.size = Pt(12.0)
                    run.font.name = 'Times New Roman'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                    if cell["row"] == 1 and cell["col"] == 0:
                        new_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER



            # for row in item["rows"]:
            #     for cell in row["cells"]:
            #         new_cell = table.cell(cell["row"], cell["col"])
            #         # 先清除段落内容
            #         for para in new_cell.paragraphs:
            #             p = para._element
            #             p.getparent().remove(p)
            #         new_cell.width = Inches(cell["width"])
            #         for para in cell["content"]:
            #             new_para = new_cell.add_paragraph()
            #             run = new_para.add_run(para["text"])
            #             run.font.size = Pt(12.0)
            #             run.font.name = 'Times New Roman'
            #             run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # 1-1
    set_cell_border(
        table.cell(0, 0),
        bottom={"sz": 12, "val": "none"},  # 底边无边框
    )
    # 1-2
    set_cell_border(
        table.cell(0, 1),
        end={"sz": 12, "val": "none"},
    )
    # 1-3
    set_cell_border(
        table.cell(0, 2),
        start={"sz": 12, "val": "none"},  # 底边无边框
    )
    # 1-4
    set_cell_border(
        table.cell(0, 3),
        end={"sz": 12, "val": "none"},
    )
    # 1-5
    set_cell_border(
        table.cell(0, 4),
        start={"sz": 12, "val": "none"},
    )
    # 2-1
    set_cell_border(
        table.cell(1, 0),
        top={"sz": 12, "val": "none"},
    )
    # 2-2
    set_cell_border(
        table.cell(1, 1),
        end={"sz": 12, "val": "none"},
    )
    # 2-3
    set_cell_border(
        table.cell(1, 2),
        start={"sz": 12, "val": "none"},
    )
    # 2-4
    set_cell_border(
        table.cell(1, 3),
        end={"sz": 12, "val": "none"},
    )
    # 2-5
    set_cell_border(
        table.cell(1, 4),
        start={"sz": 12, "val": "none"},
    )


def apply_footer_format(doc, footer_data):
    """应用页脚，生成的格式"""
    # 设置奇偶页是否相同
    doc.settings.odd_and_even_pages_header_footer = False
    # 创建节
    section = doc.sections[0]
    # 设置首页页脚不同
    section.different_first_page_header_footer = True

    # 首页页脚
    # 先清除内容
    for para in section.first_page_footer.paragraphs:
        p = para._element
        p.getparent().remove(p)
    # section.first_page_footer.add_paragraph().add_run()

    # 非首页页脚
    # 先清除内容
    for para in section.footer.paragraphs:
        p = para._element
        p.getparent().remove(p)
    # 第1行
    para1 = section.footer.add_paragraph()
    para1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    para1.paragraph_format.space_after = 0
    run1 = para1.add_run(FOOTER_FORMAT["para1"])
    run1.font.size = Pt(10)
    run1.font.name = 'Times New Roman'  # 设置西文字体
    run1.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 设置中文字体

    # 第2行
    para2 = section.footer.add_paragraph()
    para2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    para2.paragraph_format.space_after = 0
    run2 = para2.add_run(FOOTER_FORMAT["para2"])
    run2.font.size = Pt(9)
    run2.font.name = 'Times New Roman'  # 设置西文字体
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 设置中文字体
    run2.font.color.rgb = RGBColor(0xff, 0x00, 0x00)

    # 第3行
    para3 = section.footer.add_paragraph()
    para3.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    para3.paragraph_format.space_after = 0
    run3 = para3.add_run(FOOTER_FORMAT["para3"])
    run3.font.size = Pt(9)
    run3.font.name = 'Times New Roman'  # 设置西文字体
    run3.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 设置中文字体
    run3.font.color.rgb = RGBColor(0xff, 0x00, 0x00)

    # 第4行
    para4 = section.footer.add_paragraph()
    para4.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    para4.paragraph_format.space_after = 0
    run4 = para4.add_run(FOOTER_FORMAT["para4"])
    run4.font.size = Pt(9)
    run4.font.name = 'Times New Roman'  # 设置西文字体
    run4.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 设置中文字体
    run4.font.color.rgb = RGBColor(0xff, 0x00, 0x00)


def apply_paragraph_format(paragraph, format_info):
    """应用段落格式"""
    # 使用安全的get方法访问字典，避免KeyError
    if format_info.get('alignment'):
        alignment_map = {
            'CENTER': WD_ALIGN_PARAGRAPH.CENTER,
            'CENTER (1)': WD_ALIGN_PARAGRAPH.CENTER,
            'RIGHT': WD_ALIGN_PARAGRAPH.RIGHT,
            'JUSTIFY': WD_ALIGN_PARAGRAPH.JUSTIFY,
            'LEFT': WD_ALIGN_PARAGRAPH.LEFT
        }
        alignment = format_info['alignment'].split('.')[-1] if '.' in format_info['alignment'] else format_info[
            'alignment']
        paragraph.alignment = alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)

    # 使用get方法安全访问，如果键不存在则返回None
    space_before = format_info.get('space_before')
    space_after = format_info.get('space_after')
    line_spacing = format_info.get('line_spacing')
    first_line_indent = format_info.get('first_line_indent')
    left_indent = format_info.get('left_indent')

    # if space_before is not None:
    #     paragraph.paragraph_format.space_before = space_before
    if space_before:
        paragraph.paragraph_format.space_before = space_before

    if space_after:
        paragraph.paragraph_format.space_after = space_after
    else:
        paragraph.paragraph_format.space_after = 0  # 段落后间距None会默认设置为10，需设置为0

    if line_spacing:
        paragraph.paragraph_format.line_spacing = line_spacing
    if first_line_indent:
        paragraph.paragraph_format.first_line_indent = first_line_indent
    if left_indent:
        paragraph.paragraph_format.left_indent = left_indent


def apply_table_format(table, table_format_info):
    """
    处理表格样式
    """
    # 自动调整列宽：禁用
    table.autofit = False

    # 表格对齐方式，非内容对齐方式
    table.alignment = table_format_info['table_alignment']

    # 单元格样式
    for row in table_format_info['rows']:
        for cell in row['cells']:
            row = cell['row']
            col = cell['col']
            # 设置表格宽度
            if hasattr(cell, 'width'):
                table.cell(row, col).width = Inches(cell['width'])

            # 合并单元格
            if cell['is_merge_start']:
                row = cell['row']
                col = cell['col']
                # 处理水平合并（gridSpan）
                if cell['grid_span'] > 1:
                    # 合并水平单元格
                    start_cell = table.rows[row].cells[col]
                    end_cell = table.rows[row].cells[col + cell['grid_span'] - 1]
                    start_cell.merge(end_cell)
    # if table_format_info['flag'] == 'top_title':
    #     table.autofit = True


def apply_run_format(run, run_data):
    """应用运行文本格式"""
    # 使用get方法安全访问
    run.bold = run_data.get('bold') if run_data.get('bold') else False
    run.italic = run_data.get('italic') if run_data.get('italic') else False
    run.underline = run_data.get('underline') if run_data.get('underline') else False

    font_size = run_data.get('font_size')
    if font_size:
        run.font.size = Pt(font_size)

    font_name = run_data.get('font_name')
    if font_name:
        run.font.name = font_name
        run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    else:
        # 无font_name时采用默认字体
        if run.text.strip():
            run.font.name = 'Times New Roman'  # 设置西文字体
            run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 设置中文字体
        else:
            run.font.name = u'宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 设置中文字体

    font_color = run_data.get('font_color')
    if font_color:
        try:
            run.font.color.rgb = RGBColor(*font_color)
        except:
            pass


def add_content(block, data):
    """对块（doc、section、header、footer）中添加内容"""
    for index, item in enumerate(data):
        if item['type'] == 'paragraph':
            # 创建新段落
            paragraph = block.add_paragraph()
            # 应用段落格式
            apply_paragraph_format(paragraph, item['para_format'])
            # 应用运行文本、格式
            for run_data in item['runs']:
                run = paragraph.add_run(run_data['text'])
                apply_run_format(run, run_data)
        # if item['type'] == 'table':
        #     # 创建表格
        #     table = block.add_table(rows=len(item['rows']), cols=item['cols'])
        #     table.style = 'Table Grid'
        #
        #     # 设置表格样式
        #     apply_table_format(table, item)
        #
        #     # 应用表格内容
        #     for row_idx, row_data in enumerate(item['rows']):
        #         for cell_idx, cell_data in enumerate(row_data['cells']):
        #             if (cell_data['grid_span'] > 1 and cell_data['is_merge_start']) or cell_data['grid_span'] == 1:
        #                 cell = table.rows[row_idx].cells[cell_idx]
        #
        #                 # 清空默认段落
        #                 for paragraph in cell.paragraphs:
        #                     p = paragraph._element
        #                     p.getparent().remove(p)
        #
        #                 # 添加内容到单元格
        #                 if cell_data['paragraphs']:
        #                     for para_data in cell_data['paragraphs']:
        #                         cell_para = cell.add_paragraph()
        #                         apply_paragraph_format(cell_para, para_data['para_format'])
        #
        #                         for run_data in para_data['runs']:
        #                             run = cell_para.add_run(run_data['text'])
        #                             apply_run_format(run, run_data)
        #                 else:
        #                     # 如果没有详细的段落信息，只添加文本
        #                     cell.text = cell_data['text']


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")

    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f'w:{edge}'
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn(f'w:{key}'), str(edge_data[key]))

def main(data):
    current_time = datetime.datetime.now().strftime('%y%m%d%H%M%S')
    file_path = r"F:\Code\Project\tools\data\temp"
    file = os.path.join(file_path, f'test_{current_time}.docx')

    doc = Document()
    apply_footer_format(doc, data)
    apply_header_format(doc, data)

    i = 1
    while i < 30:
        doc.add_paragraph().add_run(f"这是第{i}段")
        i += 1

    doc.save(file)
    print(f"{file} was saved successfully.")


# if __name__ == '__main__':
    # doc.settings.odd_and_even_pages_header_footer = True
    #
    # for section in doc.sections:
    #     section.different_first_page_header_footer = True
    #     first_footer = section.first_page_footer
    #     first_footer.paragraphs[0].add_run("这是首页！")
    #     # first_footer.paragraphs[0].font.bold = True
    #     # first_footer.paragraphs[0].add_run().font.size = Pt(36.0)
    #     first_footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # file = r"F:\Code\Project\tools\data\data.json"
    # with open(file, 'r', encoding='utf-8') as f:
    #     data = json.load(f)
    # main(data)
