#!/usr/bin/env python3
# -*- coding: utf-8 -*- 
"""
@Project : tools
@File    : main.py
@Author  : Shawn
@Date    : 2025/9/25 11:13 
@Info    : Description of this file
"""

import datetime
import os
import sys

from docx import Document

from src.view.production.cm_sop_translate.auth import check_license
from src.view.production.cm_sop_translate.doc_process import doc_to_docx, DocumentContent, set_paper_size_format, \
    add_content, add_cover_translation, add_paragraph_translation, add_cover, add_table_translation, \
    add_header_translation, add_footer_translation
from src.view.production.cm_sop_translate.template import apply_footer_template, apply_header_template, \
    apply_header_format, apply_footer_format, apply_template
from src.view.production.cm_sop_translate.translator import Translator


def text_translate(language):
    original_text = input("请输入待翻译内容：\n").strip()
    translator = Translator()
    for lang in language:
        translator.translate(original_text, language=lang)
        # print(f"{lang}: {translator.translate(original_text, language=lang)}")


def docx_translate(language):
    # 设置输入文件夹路径
    input_folder = input("请输入待翻译的文件目录：\n")

    # 设置输出文件夹路径（二级文件夹）
    output_folder = os.path.join(input_folder, "translate_output")
    os.makedirs(output_folder, exist_ok=True)

    # 初始化翻译器
    translator = Translator()

    # 遍历文件夹中的所有文件
    for filename in os.listdir(input_folder):
        print(f"正在处理文件: {filename}")

        # 检查文件是否为Word文档（.docx格式）,若是.doc文件则转为.docx
        if filename.endswith('.doc') and not filename.startswith('~$'):  # 忽略临时文件
            doc_path = os.path.join(input_folder, filename)
            filename = os.path.basename(doc_to_docx(doc_path))

        if filename.endswith('.docx') and not filename.startswith('~$'):
            input_file = os.path.join(input_folder, filename)

            # 生成输出文件名（原文件名+时间）
            current_time = datetime.datetime.now().strftime('%y%m%d%H%M%S')
            file_base_name = os.path.splitext(filename)[0]  # 去掉扩展名
            output_filename = f"{file_base_name}_translate_{current_time}.docx"
            output_file = os.path.join(output_folder, output_filename)

            print(f"********************start at {current_time}********************")
            # 1. 读取原文档内容
            doc = Document(input_file)
            new_doc = DocumentContent(doc)
            # 正文主体
            content_data = new_doc.get_content(doc)
            # 页眉、页脚
            header_data = new_doc.get_header_content(doc)
            footer_data = new_doc.get_footer_content(doc)
            # picture&shape
            pictures = new_doc.extract_pics(doc)
            shapes = new_doc.extract_shapes(doc)

            # 2. 标注关键信息：大标题、头信息，同时修改content_data内容
            after_title = new_doc.flag_title(content_data)
            after_preamble = new_doc.flag_preamble(content_data)
            after_approve = new_doc.flag_approveTable(content_data)

            # 3. 拆分封面和正文
            data = new_doc.split_cover_body_data(after_approve)

            # 4. 翻译
            # ① 翻译封面
            translated_cover_data = add_cover_translation(data['cover_data'], translator, language)
            # ② 翻译正文内容
            after_para = add_paragraph_translation(data['body_data'], translator, language)
            translated_body_data = add_table_translation(after_para, translator, language)
            # ③ 翻译页眉、页脚
            translated_header_data = add_header_translation(header_data, translator, language)
            translated_footer_data = add_footer_translation(footer_data, translator, language)

            # 5. 处理内容
            formatted_content = apply_template(body_data=translated_body_data, header_data=translated_header_data,
                                               footer_data=translated_footer_data, cover_data=translated_cover_data)

            # 6. 创建新文档
            create_new_document(formatted_content, output_file)

            # 7. 记录此文件存在图片或者形状
            log_path = os.path.join(output_folder, "log.txt")
            log_pic = ""
            log_shape = ""
            current_time = datetime.datetime.now().strftime('%y-%m-%d %H:%M:%S')
            if len(pictures) > 0:
                log_pic = f"{current_time}: {filename}存在{len(pictures)}张图片！\n"
            if len(shapes) > 0:
                log_shape = f"{current_time}: {filename}存在{len(shapes)}个形状！\n"
            if (log_pic+log_shape).strip():
                with open(log_path, "a", encoding="utf-8") as f:
                    f.write(log_pic+log_shape)
            print(f"  已完成翻译: {output_filename}")

    print(f"所有文件处理完成！输出目录: {output_folder}")


def create_new_document(data, output_path):
    """
    根据记录的内容和格式生成新的Word文档
    """
    doc = Document()
    set_paper_size_format(doc)

    try:
        apply_header_format(doc, data['header'])
        apply_footer_format(doc, data['footer'])
        add_cover(doc, data['cover'])
        add_content(doc, data['body'])
    except Exception as e:
        print(f"生成文件失败，异常信息：{e}")
    finally:
        # 保存文档
        doc.save(output_path)
        print(f"新文档已保存到: {output_path}")


if __name__ == '__main__':
    print("\n" + "=" * 100)
    print("\x20" * 45 + f"文档翻译系统")
    print("=" * 100)
    print("注意：仅支持运行于windows系统！！！")
    # 首先进行登录验证
    # if not login():
    #     exit()

    # 可选：进行许可证检查
    if not check_license():
        exit()

    while True:
        try:
            print("请选择使用方式(请输入序号)：\n"
                  "1. 文本翻译\n"
                  "2. 文档翻译\n"
                  "3. 退出")
            option = input().strip()
            language = ['英语', '越南语']
            if option == '1':
                print(f"当前目标语言为{language}")
                text_translate(language)
            elif option == '2':
                print(f"当前目标语言为{language}")
                docx_translate(language)
            elif option == '3':
                print("\n感谢使用！程序即将退出...")
                input("按回车键关闭窗口...")
                sys.exit(0)
            else:
                print("输入错误，请重新输入")

        except Exception as e:
            print(f"处理过程中出现错误: {e}")
        finally:
            current_time = datetime.datetime.now().strftime('%y%m%d%H%M%S')
            print(f"\n********************{current_time} end**********************")
